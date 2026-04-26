"""Tests for cli.py module."""

from __future__ import annotations

import json
import tempfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest
from typer.testing import CliRunner

from app.cli import app, score_command

try:
    import pymupdf
except Exception:  # noqa: BLE001 - environment-dependent optional dependency
    pymupdf = None

runner = CliRunner()


# ============================================================================
# Fixtures
# ============================================================================
@pytest.fixture
def sample_shigong_path():
    """Return path to sample shigong file."""
    return str(Path(__file__).parent.parent / "sample_shigong.txt")


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test outputs."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def mock_rules_report():
    """Create a mock ScoreReport for testing."""
    mock_report = MagicMock()
    mock_report.total_score = 75.0
    mock_report.judge_mode = "rules"
    mock_report.judge_source = "rules_engine"
    mock_report.spark_called = False
    mock_report.model_dump.return_value = {
        "total_score": 75.0,
        "judge_mode": "rules",
        "judge_source": "rules_engine",
        "spark_called": False,
        "dimensions": [],
        "penalties": [],
        "suggestions": [],
    }
    return mock_report


# ============================================================================
# Tests for CLI app structure
# ============================================================================
class TestCliStructure:
    """Tests for CLI app structure."""

    def test_app_exists(self):
        """Test that the main app exists."""
        assert app is not None

    def test_app_has_score_command(self):
        """Test that score command is registered."""
        # Check that 'score' is a registered command/typer
        command_names = [cmd.name for cmd in app.registered_groups]
        assert "score" in command_names

    def test_app_has_warmup_command(self):
        """Test that warmup command is registered."""
        command_names = [cmd.name for cmd in app.registered_groups]
        assert "warmup" in command_names

    def test_app_has_agents_command(self):
        """Test that agents command is registered."""
        command_names = [cmd.name for cmd in app.registered_groups]
        assert "agents" in command_names

    def test_help_option(self):
        """Test --help option shows usage."""
        result = runner.invoke(app, ["--help"])
        assert result.exit_code == 0
        assert "施工组织设计评标评分" in result.stdout

    def test_score_help(self):
        """Test score --help shows usage."""
        result = runner.invoke(app, ["score", "--help"])
        assert result.exit_code == 0
        assert "施组文本" in result.stdout or "input" in result.stdout.lower()

    def test_warmup_help(self):
        """Test warmup --help shows usage."""
        result = runner.invoke(app, ["warmup", "--help"])
        assert result.exit_code == 0
        assert "预热" in result.stdout or "warmup" in result.stdout.lower()
        assert "input" in result.stdout.lower() or "--input" in result.stdout


# ============================================================================
# Tests for warmup command
# ============================================================================
class TestWarmupCommand:
    """Tests for warmup command."""

    def test_warmup_nonexistent_file(self, temp_dir):
        """Test warmup with nonexistent file exits with error."""
        result = runner.invoke(app, ["warmup", "-i", str(Path(temp_dir) / "nonexistent.txt")])
        assert result.exit_code == 1
        assert "不存在" in result.stdout or "not found" in result.stdout.lower()

    @patch("app.cli.score_text")
    @patch("app.cli.load_config")
    def test_warmup_from_txt_file(self, mock_load_config, mock_score_text, temp_dir):
        """Test warmup from a .txt file (one line per item) with mocked scorer."""
        mock_load_config.return_value = MagicMock(
            rubric={},
            lexicon={},
        )
        mock_report = MagicMock()
        mock_report.model_dump.return_value = {"total_score": 80.0, "dimensions": []}
        mock_score_text.return_value = mock_report

        txt_path = Path(temp_dir) / "items.txt"
        txt_path.write_text("短文本一行\n另一行内容\n", encoding="utf-8")

        result = runner.invoke(app, ["warmup", "-i", str(txt_path)])
        assert result.exit_code == 0
        assert "预热完成" in result.stdout
        assert "共 2 条" in result.stdout or "2" in result.stdout
        assert mock_score_text.call_count == 2


class TestAgentsCommand:
    def test_agents_list_outputs_registered_specs(self):
        services = SimpleNamespace(
            agents=SimpleNamespace(
                list_agents=MagicMock(return_value=[{"agent_name": "ops-triage"}])
            )
        )
        with patch("app.cli.get_application_services", return_value=services):
            result = runner.invoke(app, ["agents", "list"])

        assert result.exit_code == 0
        payload = json.loads(result.stdout)
        assert payload[0]["agent_name"] == "ops-triage"

    def test_agents_dry_run_prints_result_payload(self):
        run_result = SimpleNamespace(
            model_dump=lambda mode="json": {
                "audit": {"status": "success"},
                "output": {"overall_status": "warn"},
            }
        )
        services = SimpleNamespace(
            agents=SimpleNamespace(dry_run=MagicMock(return_value=run_result))
        )
        with patch("app.cli.get_application_services", return_value=services):
            result = runner.invoke(app, ["agents", "dry-run", "--agent", "ops-triage"])

        assert result.exit_code == 0
        payload = json.loads(result.stdout)
        assert payload["audit"]["status"] == "success"
        assert payload["output"]["overall_status"] == "warn"


# ============================================================================
# Tests for score command with rules mode
# ============================================================================
class TestScoreRulesMode:
    """Tests for score command in rules mode."""

    def test_score_rules_mode_basic(self, sample_shigong_path):
        """Test basic scoring in rules mode."""
        result = runner.invoke(app, ["score", "-i", sample_shigong_path, "--mode", "rules"])
        assert result.exit_code == 0
        # Output should be valid JSON
        output = json.loads(result.stdout)
        assert "judge_mode" in output or "total_score" in output

    def test_score_rules_mode_sets_attributes(self, sample_shigong_path):
        """Test that rules mode sets correct attributes."""
        result = runner.invoke(app, ["score", "-i", sample_shigong_path, "--mode", "rules"])
        assert result.exit_code == 0
        output = json.loads(result.stdout)
        assert output.get("judge_mode") == "rules"
        assert output.get("judge_source") == "rules_engine"
        assert output.get("spark_called") is False

    def test_score_with_output_file(self, sample_shigong_path, temp_dir):
        """Test outputting score to file."""
        out_path = Path(temp_dir) / "output.json"
        result = runner.invoke(
            app, ["score", "-i", sample_shigong_path, "--mode", "rules", "-o", str(out_path)]
        )
        assert result.exit_code == 0
        assert out_path.exists()
        content = json.loads(out_path.read_text(encoding="utf-8"))
        assert "judge_mode" in content or "total_score" in content

    def test_score_with_summary(self, sample_shigong_path):
        """Test score with summary output."""
        result = runner.invoke(
            app, ["score", "-i", sample_shigong_path, "--mode", "rules", "--summary"]
        )
        assert result.exit_code == 0
        # Should contain both JSON and report text
        assert "judge_mode" in result.stdout or "total_score" in result.stdout

    def test_score_with_summary_out(self, sample_shigong_path, temp_dir):
        """Test score with summary output to file."""
        summary_path = Path(temp_dir) / "summary.txt"
        result = runner.invoke(
            app,
            [
                "score",
                "-i",
                sample_shigong_path,
                "--mode",
                "rules",
                "--summary",
                "--summary-out",
                str(summary_path),
            ],
        )
        assert result.exit_code == 0
        assert summary_path.exists()
        assert "已生成报告" in result.stdout

    def test_score_with_docx_out(self, sample_shigong_path, temp_dir):
        """Test score with DOCX output."""
        docx_path = Path(temp_dir) / "report.docx"
        result = runner.invoke(
            app,
            ["score", "-i", sample_shigong_path, "--mode", "rules", "--docx-out", str(docx_path)],
        )
        assert result.exit_code == 0
        assert docx_path.exists()
        assert "已生成 DOCX 报告" in result.stdout


# ============================================================================
# Tests for score command with spark mode
# ============================================================================
class TestScoreSparkMode:
    """Tests for score command in spark mode."""

    @patch("app.cli.run_spark_judge")
    def test_spark_mode_with_api_success(self, mock_spark, sample_shigong_path):
        """Test spark mode when API call succeeds."""
        mock_spark.return_value = {
            "called_spark_api": True,
            "overall": {"total_score_0_100": 80.0},
            "prompt_version": "test_v1",
        }
        result = runner.invoke(app, ["score", "-i", sample_shigong_path, "--mode", "spark"])
        assert result.exit_code == 0
        output = json.loads(result.stdout)
        assert output.get("judge_mode") == "openai"
        assert output.get("judge_source") == "openai_api"

    @patch("app.cli.run_spark_judge")
    def test_spark_mode_returns_interrupted_payload_when_api_fails(
        self, mock_spark, sample_shigong_path
    ):
        """Test spark mode surfaces interruption instead of silently falling back."""
        mock_spark.return_value = {
            "called_spark_api": False,
            "processing_interrupted": True,
            "message": "计算中断异常，请重试",
            "reason": "API unavailable",
        }
        result = runner.invoke(app, ["score", "-i", sample_shigong_path, "--mode", "spark"])
        assert result.exit_code == 0
        output = json.loads(result.stdout)
        assert output.get("judge_mode") == "openai_interrupted"
        assert output.get("judge_source") == "openai_api"
        assert output.get("spark_called") is False
        assert output.get("processing_interrupted") is True
        assert "fallback_reason" in output


# ============================================================================
# Tests for score command with hybrid mode
# ============================================================================
class TestScoreHybridMode:
    """Tests for score command in hybrid mode."""

    @patch("app.cli.run_spark_judge")
    def test_hybrid_mode_with_api_success(self, mock_spark, sample_shigong_path):
        """Test hybrid mode when API call succeeds."""
        mock_spark.return_value = {
            "called_spark_api": True,
            "overall": {"total_score_0_100": 82.0},
            "prompt_version": "test_v1",
        }
        result = runner.invoke(app, ["score", "-i", sample_shigong_path, "--mode", "hybrid"])
        assert result.exit_code == 0
        output = json.loads(result.stdout)
        assert output.get("judge_mode") == "hybrid"
        assert output.get("judge_source") == "openai_api"
        assert output.get("spark_called") is True
        assert "base_rules_score" in output
        assert "llm_adjustment" in output
        assert "final_total_score" in output

    @patch("app.cli.run_spark_judge")
    def test_hybrid_mode_adjustment_capped(self, mock_spark, sample_shigong_path):
        """Test hybrid mode caps adjustment within max_adjustment."""
        mock_spark.return_value = {
            "called_spark_api": True,
            "overall": {"total_score_0_100": 200.0},  # Extreme value
            "prompt_version": "test_v1",
        }
        result = runner.invoke(app, ["score", "-i", sample_shigong_path, "--mode", "hybrid"])
        assert result.exit_code == 0
        output = json.loads(result.stdout)
        # Adjustment should be capped (default max_adjustment is 10)
        assert abs(output.get("llm_adjustment", 0)) <= 10.0

    @patch("app.cli.run_spark_judge")
    def test_hybrid_mode_returns_interrupted_payload_when_api_fails(
        self, mock_spark, sample_shigong_path
    ):
        """Test hybrid mode surfaces interruption instead of silently falling back."""
        mock_spark.return_value = {
            "called_spark_api": False,
            "processing_interrupted": True,
            "message": "计算中断异常，请重试",
            "reason": "Connection timeout",
        }
        result = runner.invoke(app, ["score", "-i", sample_shigong_path, "--mode", "hybrid"])
        assert result.exit_code == 0
        output = json.loads(result.stdout)
        assert output.get("judge_mode") == "hybrid_interrupted"
        assert output.get("judge_source") == "openai_api"
        assert output.get("spark_called") is False
        assert output.get("processing_interrupted") is True
        assert "fallback_reason" in output


# ============================================================================
# Tests for error handling
# ============================================================================
class TestErrorHandling:
    """Tests for CLI error handling."""

    def test_invalid_mode(self, sample_shigong_path):
        """Test error on invalid mode."""
        result = runner.invoke(app, ["score", "-i", sample_shigong_path, "--mode", "invalid"])
        assert result.exit_code != 0

    def test_missing_input(self):
        """Test error when input file is missing."""
        result = runner.invoke(app, ["score", "--mode", "rules"])
        assert result.exit_code != 0

    def test_nonexistent_input_file(self):
        """Test error when input file doesn't exist."""
        result = runner.invoke(app, ["score", "-i", "/nonexistent/path.txt", "--mode", "rules"])
        assert result.exit_code != 0


# ============================================================================
# Tests for main entry point
# ============================================================================
class TestMainEntryPoint:
    """Tests for main entry point."""

    def test_main_module_invocable(self):
        """Test that app() can be called as main entry point."""
        # Just verify the app is callable
        assert callable(app)

    def test_score_command_callable(self):
        """Test that score_command function exists and is callable."""
        assert callable(score_command)


# ============================================================================
# Tests for combined options
# ============================================================================
class TestCombinedOptions:
    """Tests for combined CLI options."""

    def test_all_output_options(self, sample_shigong_path, temp_dir):
        """Test using multiple output options together."""
        out_path = Path(temp_dir) / "output.json"
        summary_path = Path(temp_dir) / "summary.txt"
        docx_path = Path(temp_dir) / "report.docx"
        result = runner.invoke(
            app,
            [
                "score",
                "-i",
                sample_shigong_path,
                "--mode",
                "rules",
                "-o",
                str(out_path),
                "--summary",
                "--summary-out",
                str(summary_path),
                "--docx-out",
                str(docx_path),
            ],
        )
        assert result.exit_code == 0
        assert out_path.exists()
        assert summary_path.exists()
        assert docx_path.exists()

    def test_output_file_content_matches_stdout(self, sample_shigong_path, temp_dir):
        """Test that output file content matches what's printed to stdout."""
        out_path = Path(temp_dir) / "output.json"
        result = runner.invoke(
            app, ["score", "-i", sample_shigong_path, "--mode", "rules", "-o", str(out_path)]
        )
        assert result.exit_code == 0
        file_content = json.loads(out_path.read_text(encoding="utf-8"))
        # First part of stdout should be JSON
        stdout_json = json.loads(
            result.stdout.strip().split("\n{")[0] + "{"
            if "\n{" in result.stdout
            else result.stdout.strip()
        )
        # Both should have same total_score
        assert file_content.get("total_score") == stdout_json.get("total_score")


# ============================================================================
# Tests for batch command
# ============================================================================
class TestBatchCommand:
    """Tests for batch command functionality."""

    def test_batch_help(self):
        """Test batch --help shows usage."""
        result = runner.invoke(app, ["batch", "--help"])
        assert result.exit_code == 0
        assert "批量" in result.stdout or "input" in result.stdout.lower()

    def test_batch_single_file(self, sample_shigong_path, temp_dir):
        """Test batch processing a single file."""
        result = runner.invoke(app, ["batch", "-i", sample_shigong_path, "-o", temp_dir])
        assert result.exit_code == 0
        assert "1 个文件" in result.stdout
        assert "成功" in result.stdout
        # Check output files exist
        out_path = Path(temp_dir)
        json_files = list(out_path.glob("*_report.json"))
        assert len(json_files) == 1
        # Check summary file
        summary_path = out_path / "_batch_summary.json"
        assert summary_path.exists()
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        assert summary["success_count"] == 1

    def test_batch_multiple_files(self, sample_shigong_path, temp_dir):
        """Test batch processing multiple files."""
        # Create a second test file
        test_file2 = Path(temp_dir) / "test_input2.txt"
        test_file2.write_text("测试施工组织设计文本", encoding="utf-8")
        result = runner.invoke(
            app, ["batch", "-i", sample_shigong_path, "-i", str(test_file2), "-o", temp_dir]
        )
        assert result.exit_code == 0
        assert "2 个文件" in result.stdout
        # Check output files
        out_path = Path(temp_dir)
        json_files = list(out_path.glob("*_report.json"))
        assert len(json_files) == 2

    def test_batch_with_docx(self, sample_shigong_path, temp_dir):
        """Test batch processing with DOCX output."""
        result = runner.invoke(app, ["batch", "-i", sample_shigong_path, "-o", temp_dir, "--docx"])
        assert result.exit_code == 0
        out_path = Path(temp_dir)
        docx_files = list(out_path.glob("*_report.docx"))
        assert len(docx_files) == 1

    def test_batch_directory_input(self, temp_dir):
        """Test batch processing a directory of files."""
        # Create input directory with test files
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        for i in range(3):
            (input_dir / f"test_{i}.txt").write_text(f"测试文档 {i}", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"
        result = runner.invoke(
            app, ["batch", "-i", str(input_dir), "-o", str(output_dir), "--pattern", "*.txt"]
        )
        assert result.exit_code == 0
        assert "3 个文件" in result.stdout
        json_files = list(output_dir.glob("*_report.json"))
        assert len(json_files) == 3

    def test_batch_no_files_found(self, temp_dir):
        """Test batch with no matching files."""
        empty_dir = Path(temp_dir) / "empty"
        empty_dir.mkdir()
        result = runner.invoke(app, ["batch", "-i", str(empty_dir), "-o", temp_dir])
        assert result.exit_code == 1
        assert "没有找到" in result.stdout

    def test_batch_nonexistent_path(self, temp_dir):
        """Test batch with nonexistent path shows warning."""
        result = runner.invoke(app, ["batch", "-i", "/nonexistent/path.txt", "-o", temp_dir])
        # Should warn and fail (no files)
        assert "警告" in result.stdout or "没有找到" in result.stdout

    def test_batch_summary_content(self, sample_shigong_path, temp_dir):
        """Test batch summary contains expected fields."""
        result = runner.invoke(app, ["batch", "-i", sample_shigong_path, "-o", temp_dir])
        assert result.exit_code == 0
        summary_path = Path(temp_dir) / "_batch_summary.json"
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        assert "total_files" in summary
        assert "success_count" in summary
        assert "error_count" in summary
        assert "results" in summary
        assert len(summary["results"]) == 1
        assert summary["results"][0]["status"] == "success"
        assert "total_score" in summary["results"][0]


# ============================================================================
# Tests for batch parallel processing
# ============================================================================
class TestBatchParallel:
    """Tests for batch command parallel processing functionality."""

    def test_batch_parallel_single_file(self, sample_shigong_path, temp_dir):
        """Test parallel batch with single file (should work same as serial)."""
        result = runner.invoke(
            app, ["batch", "-i", sample_shigong_path, "-o", temp_dir, "--workers", "2"]
        )
        assert result.exit_code == 0
        # Single file, effective workers = 1, should not show parallel mode
        assert "成功" in result.stdout
        summary_path = Path(temp_dir) / "_batch_summary.json"
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        assert summary["success_count"] == 1

    def test_batch_parallel_multiple_files(self, temp_dir):
        """Test parallel batch with multiple files."""
        # Create input directory with test files
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        for i in range(4):
            (input_dir / f"test_{i}.txt").write_text(
                f"测试施工组织设计文档 {i}\n工期计划：30天", encoding="utf-8"
            )
        output_dir = Path(temp_dir) / "outputs"
        result = runner.invoke(
            app,
            ["batch", "-i", str(input_dir), "-o", str(output_dir), "--workers", "4"],
        )
        assert result.exit_code == 0
        assert "4 个文件" in result.stdout
        assert "并行模式：4 个工作线程" in result.stdout
        json_files = list(output_dir.glob("*_report.json"))
        assert len(json_files) == 4
        summary_path = output_dir / "_batch_summary.json"
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        assert summary["success_count"] == 4

    def test_batch_parallel_with_docx(self, temp_dir):
        """Test parallel batch with DOCX output."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        for i in range(3):
            (input_dir / f"test_{i}.txt").write_text(
                f"测试文档 {i}\n工期：{i * 10 + 10}天", encoding="utf-8"
            )
        output_dir = Path(temp_dir) / "outputs"
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--workers",
                "3",
                "--docx",
            ],
        )
        assert result.exit_code == 0
        assert "并行模式：3 个工作线程" in result.stdout
        docx_files = list(output_dir.glob("*_report.docx"))
        assert len(docx_files) == 3

    def test_batch_workers_capped_by_file_count(self, temp_dir):
        """Test that workers are capped by file count."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        for i in range(2):
            (input_dir / f"test_{i}.txt").write_text(f"测试文档 {i}", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"
        # Request 10 workers but only 2 files
        result = runner.invoke(
            app,
            ["batch", "-i", str(input_dir), "-o", str(output_dir), "--workers", "10"],
        )
        assert result.exit_code == 0
        assert "2 个文件" in result.stdout
        # Should cap workers to 2
        assert "并行模式：2 个工作线程" in result.stdout

    def test_batch_serial_with_workers_1(self, sample_shigong_path, temp_dir):
        """Test that workers=1 uses serial mode."""
        # Create a second test file
        test_file2 = Path(temp_dir) / "test_input2.txt"
        test_file2.write_text("测试施工组织设计文本", encoding="utf-8")
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                sample_shigong_path,
                "-i",
                str(test_file2),
                "-o",
                temp_dir,
                "--workers",
                "1",
            ],
        )
        assert result.exit_code == 0
        # Should NOT show parallel mode message
        assert "并行模式" not in result.stdout
        # Should show serial processing format
        assert "处理：" in result.stdout

    def test_batch_parallel_error_handling(self, temp_dir):
        """Test parallel batch handles errors gracefully."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        # Create valid files
        (input_dir / "valid1.txt").write_text("有效文档1", encoding="utf-8")
        (input_dir / "valid2.txt").write_text("有效文档2", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"
        result = runner.invoke(
            app,
            ["batch", "-i", str(input_dir), "-o", str(output_dir), "--workers", "2"],
        )
        assert result.exit_code == 0
        summary_path = output_dir / "_batch_summary.json"
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        # Both should succeed
        assert summary["success_count"] == 2


# ============================================================================
# Tests for DOCX input support
# ============================================================================
class TestDocxInputSupport:
    """Tests for DOCX input file support."""

    def test_read_input_file_txt(self, temp_dir):
        """Test read_input_file with .txt file."""
        from app.cli import read_input_file

        txt_path = Path(temp_dir) / "test.txt"
        txt_path.write_text("这是测试文本内容", encoding="utf-8")
        result = read_input_file(txt_path)
        assert result == "这是测试文本内容"

    def test_read_input_file_docx(self, temp_dir):
        """Test read_input_file with .docx file."""
        from docx import Document

        from app.cli import read_input_file

        docx_path = Path(temp_dir) / "test.docx"
        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")
        doc.save(str(docx_path))

        result = read_input_file(docx_path)
        assert "第一段内容" in result
        assert "第二段内容" in result

    def test_read_input_file_pdf(self, temp_dir):
        """Test read_input_file with .pdf file."""
        from app.cli import read_input_file

        if pymupdf is None:
            pytest.skip("PyMuPDF not available in this environment")

        pdf_path = Path(temp_dir) / "test.pdf"
        # 创建真实的 PDF 文件，使用支持中文的字体
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Page1: Construction Plan", fontname="helv")
        page2 = doc.new_page()
        page2.insert_text((72, 72), "Page2: Safety Measures", fontname="helv")
        doc.save(str(pdf_path))
        doc.close()

        result = read_input_file(pdf_path)
        assert "Construction Plan" in result
        assert "Safety Measures" in result

    def test_read_input_file_unsupported_format(self, temp_dir):
        """Test read_input_file raises error for unsupported format."""
        from app.cli import read_input_file

        xyz_path = Path(temp_dir) / "test.xyz"
        xyz_path.write_text("fake file", encoding="utf-8")

        with pytest.raises(ValueError, match="不支持的文件格式"):
            read_input_file(xyz_path)

    def test_score_command_docx_input(self, temp_dir):
        """Test score command with DOCX input file."""
        from docx import Document

        # Create a sample DOCX input
        docx_path = Path(temp_dir) / "input.docx"
        doc = Document()
        doc.add_paragraph("施工组织设计")
        doc.add_paragraph("本工程为某建筑项目的施工组织设计。")
        doc.save(str(docx_path))

        output_path = Path(temp_dir) / "output.json"
        result = runner.invoke(
            app,
            ["score", "-i", str(docx_path), "-o", str(output_path)],
        )
        assert result.exit_code == 0
        assert output_path.exists()
        report = json.loads(output_path.read_text(encoding="utf-8"))
        assert "total_score" in report

    def test_score_command_pdf_input(self, temp_dir):
        """Test score command with PDF input file."""
        # Create a sample PDF input with English text (font support)
        if pymupdf is None:
            pytest.skip("PyMuPDF not available in this environment")

        pdf_path = Path(temp_dir) / "input.pdf"
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Construction Organization Design", fontname="helv")
        page.insert_text((72, 100), "This is a construction project plan.", fontname="helv")
        doc.save(str(pdf_path))
        doc.close()

        output_path = Path(temp_dir) / "output.json"
        result = runner.invoke(
            app,
            ["score", "-i", str(pdf_path), "-o", str(output_path)],
        )
        assert result.exit_code == 0
        assert output_path.exists()
        report = json.loads(output_path.read_text(encoding="utf-8"))
        assert "total_score" in report

    def test_batch_command_docx_input(self, temp_dir):
        """Test batch command with DOCX input files."""
        from docx import Document

        # Create sample DOCX inputs
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()

        for i in range(2):
            docx_path = input_dir / f"doc{i}.docx"
            doc = Document()
            doc.add_paragraph(f"施工组织设计文档 {i}")
            doc.add_paragraph("本工程为某建筑项目的施工组织设计。")
            doc.save(str(docx_path))

        output_dir = Path(temp_dir) / "outputs"
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--pattern",
                "*.docx",
            ],
        )
        assert result.exit_code == 0
        summary_path = output_dir / "_batch_summary.json"
        assert summary_path.exists()
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        assert summary["total_files"] == 2
        assert summary["success_count"] == 2

    def test_batch_command_pdf_input(self, temp_dir):
        """Test batch command with PDF input files."""
        # Create sample PDF inputs with English text (font support)
        if pymupdf is None:
            pytest.skip("PyMuPDF not available in this environment")

        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()

        for i in range(2):
            pdf_path = input_dir / f"doc{i}.pdf"
            doc = pymupdf.open()
            page = doc.new_page()
            page.insert_text((72, 72), f"Construction document {i}", fontname="helv")
            page.insert_text((72, 100), "This is a construction project plan.", fontname="helv")
            doc.save(str(pdf_path))
            doc.close()

        output_dir = Path(temp_dir) / "outputs"
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--pattern",
                "*.pdf",
            ],
        )
        assert result.exit_code == 0
        summary_path = output_dir / "_batch_summary.json"
        assert summary_path.exists()
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        assert summary["total_files"] == 2
        assert summary["success_count"] == 2

    def test_batch_command_mixed_formats(self, temp_dir):
        """Test batch command with mixed .txt and .docx files."""
        from docx import Document

        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()

        # Create one txt file
        (input_dir / "doc1.txt").write_text("施工组织设计文本文件", encoding="utf-8")

        # Process txt first
        output_dir = Path(temp_dir) / "outputs"
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir / "doc1.txt"),
                "-o",
                str(output_dir),
            ],
        )
        assert result.exit_code == 0

        # Create one docx file and process it
        docx_path = input_dir / "doc2.docx"
        doc = Document()
        doc.add_paragraph("施工组织设计 DOCX 文件")
        doc.save(str(docx_path))

        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(docx_path),
                "-o",
                str(output_dir),
            ],
        )
        assert result.exit_code == 0
        # Should have output for docx file
        assert (output_dir / "doc2_report.json").exists()


# ============================================================================
# Tests for ProcessPoolExecutor support
# ============================================================================
class TestProcessPoolExecutor:
    """Tests for ProcessPoolExecutor (--executor process) option."""

    def test_batch_executor_help_shows_option(self):
        """Test that --executor option is shown in help."""
        result = runner.invoke(app, ["batch", "--help"])
        assert result.exit_code == 0
        assert "--executor" in result.output
        assert "thread" in result.output
        assert "process" in result.output

    def test_batch_process_executor_two_files(self, temp_dir):
        """Test batch with ProcessPoolExecutor on two files (need 2+ for parallel)."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        (input_dir / "test1.txt").write_text("施工组织设计测试1", encoding="utf-8")
        (input_dir / "test2.txt").write_text("施工组织设计测试2", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"

        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--workers",
                "2",
                "--executor",
                "process",
            ],
        )
        assert result.exit_code == 0
        assert "工作进程" in result.output  # Should say "进程" not "线程"
        assert (output_dir / "test1_report.json").exists()
        assert (output_dir / "test2_report.json").exists()

    def test_batch_process_executor_multiple_files(self, temp_dir):
        """Test batch with ProcessPoolExecutor on multiple files."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        for i in range(3):
            (input_dir / f"doc{i}.txt").write_text(f"测试文档 {i}", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"

        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--workers",
                "2",
                "--executor",
                "process",
            ],
        )
        assert result.exit_code == 0
        assert "并行模式：2 个工作进程" in result.output
        summary_path = output_dir / "_batch_summary.json"
        summary = json.loads(summary_path.read_text(encoding="utf-8"))
        assert summary["total_files"] == 3
        assert summary["success_count"] == 3

    def test_batch_thread_executor_explicit(self, temp_dir):
        """Test batch with explicit --executor thread (need 2+ files for parallel)."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        (input_dir / "test1.txt").write_text("施工组织设计测试1", encoding="utf-8")
        (input_dir / "test2.txt").write_text("施工组织设计测试2", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"

        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--workers",
                "2",
                "--executor",
                "thread",
            ],
        )
        assert result.exit_code == 0
        assert "工作线程" in result.output  # Should say "线程"
        assert (output_dir / "test1_report.json").exists()
        assert (output_dir / "test2_report.json").exists()

    def test_batch_invalid_executor(self, temp_dir):
        """Test batch with invalid executor value."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        (input_dir / "test.txt").write_text("测试", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"

        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir / "test.txt"),
                "-o",
                str(output_dir),
                "--workers",
                "2",
                "--executor",
                "invalid",
            ],
        )
        assert result.exit_code == 1
        assert "executor" in result.output.lower()

    def test_batch_process_executor_with_docx(self, temp_dir):
        """Test batch ProcessPoolExecutor with DOCX output."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        for i in range(2):
            (input_dir / f"doc{i}.txt").write_text(f"施工组织设计 {i}", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"

        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--workers",
                "2",
                "--executor",
                "process",
                "--docx",
            ],
        )
        assert result.exit_code == 0
        # Check both JSON and DOCX outputs exist
        assert (output_dir / "doc0_report.json").exists()
        assert (output_dir / "doc0_report.docx").exists()
        assert (output_dir / "doc1_report.json").exists()
        assert (output_dir / "doc1_report.docx").exists()

    @patch("app.cli.ProcessPoolExecutor")
    def test_batch_process_executor_falls_back_when_process_pool_infra_unavailable(
        self,
        mock_process_executor,
        temp_dir,
    ):
        """Test batch auto-falls back to thread pool when process infra is blocked."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        (input_dir / "test1.txt").write_text("施工组织设计测试1", encoding="utf-8")
        (input_dir / "test2.txt").write_text("施工组织设计测试2", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"

        mock_process_executor.side_effect = ImportError(
            "dlopen(_posixshmem): library load denied by system policy"
        )

        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--workers",
                "2",
                "--executor",
                "process",
            ],
        )

        assert result.exit_code == 0
        assert "已自动回退为线程池" in result.output
        assert (output_dir / "test1_report.json").exists()
        assert (output_dir / "test2_report.json").exists()

    def test_batch_executor_default_is_thread(self, temp_dir):
        """Test that default executor is thread (need 2+ files for parallel)."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        (input_dir / "test1.txt").write_text("测试1", encoding="utf-8")
        (input_dir / "test2.txt").write_text("测试2", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"

        # Without --executor, should use thread by default
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--workers",
                "2",
            ],
        )
        assert result.exit_code == 0
        assert "工作线程" in result.output  # Default is thread


# ============================================================================
# Tests for CLI i18n (--locale) support
# ============================================================================
class TestCliI18n:
    """Tests for CLI internationalization support."""

    def test_score_help_shows_locale_option(self):
        """Test that --locale option is shown in score help."""
        result = runner.invoke(app, ["score", "--help"])
        assert result.exit_code == 0
        assert "--locale" in result.output or "-l" in result.output

    def test_batch_help_shows_locale_option(self):
        """Test that --locale option is shown in batch help."""
        result = runner.invoke(app, ["batch", "--help"])
        assert result.exit_code == 0
        assert "--locale" in result.output or "-l" in result.output

    def test_score_with_locale_zh(self, sample_shigong_path, temp_dir):
        """Test score command with Chinese locale."""
        docx_path = Path(temp_dir) / "report.docx"
        result = runner.invoke(
            app,
            [
                "score",
                "-i",
                sample_shigong_path,
                "--mode",
                "rules",
                "--docx-out",
                str(docx_path),
                "--locale",
                "zh",
            ],
        )
        assert result.exit_code == 0
        assert docx_path.exists()
        # Chinese output message
        assert "已生成 DOCX 报告" in result.stdout

    def test_score_with_locale_en(self, sample_shigong_path, temp_dir):
        """Test score command with English locale."""
        docx_path = Path(temp_dir) / "report.docx"
        result = runner.invoke(
            app,
            [
                "score",
                "-i",
                sample_shigong_path,
                "--mode",
                "rules",
                "--docx-out",
                str(docx_path),
                "--locale",
                "en",
            ],
        )
        assert result.exit_code == 0
        assert docx_path.exists()
        # English output message
        assert "DOCX report generated" in result.stdout

    def test_score_with_summary_locale_en(self, sample_shigong_path, temp_dir):
        """Test score with summary output in English."""
        summary_path = Path(temp_dir) / "summary.txt"
        result = runner.invoke(
            app,
            [
                "score",
                "-i",
                sample_shigong_path,
                "--mode",
                "rules",
                "--summary",
                "--summary-out",
                str(summary_path),
                "--locale",
                "en",
            ],
        )
        assert result.exit_code == 0
        assert summary_path.exists()
        # English report generated message
        assert "Report generated" in result.stdout
        # Check summary content is in English
        summary_content = summary_path.read_text(encoding="utf-8")
        assert "Qingtian" in summary_content or "Total Score" in summary_content

    def test_score_invalid_locale(self, sample_shigong_path):
        """Test score command with invalid locale."""
        result = runner.invoke(
            app,
            [
                "score",
                "-i",
                sample_shigong_path,
                "--mode",
                "rules",
                "--locale",
                "fr",  # French not supported
            ],
        )
        assert result.exit_code != 0

    def test_batch_with_locale_zh(self, sample_shigong_path, temp_dir):
        """Test batch command with Chinese locale."""
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                sample_shigong_path,
                "-o",
                temp_dir,
                "--locale",
                "zh",
            ],
        )
        assert result.exit_code == 0
        # Chinese output messages
        assert "个文件待处理" in result.stdout
        assert "成功" in result.stdout

    def test_batch_with_locale_en(self, sample_shigong_path, temp_dir):
        """Test batch command with English locale."""
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                sample_shigong_path,
                "-o",
                temp_dir,
                "--locale",
                "en",
            ],
        )
        assert result.exit_code == 0
        # English output messages
        assert "files to process" in result.stdout or "Found" in result.stdout
        assert "succeeded" in result.stdout or "Completed" in result.stdout

    def test_batch_parallel_with_locale_en(self, temp_dir):
        """Test batch parallel with English locale."""
        input_dir = Path(temp_dir) / "inputs"
        input_dir.mkdir()
        for i in range(2):
            (input_dir / f"test_{i}.txt").write_text(f"Test document {i}", encoding="utf-8")
        output_dir = Path(temp_dir) / "outputs"

        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                str(input_dir),
                "-o",
                str(output_dir),
                "--workers",
                "2",
                "--locale",
                "en",
            ],
        )
        assert result.exit_code == 0
        # English parallel mode message
        assert "worker threads" in result.stdout or "Parallel mode" in result.stdout

    def test_batch_invalid_locale(self, sample_shigong_path, temp_dir):
        """Test batch command with invalid locale."""
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                sample_shigong_path,
                "-o",
                temp_dir,
                "--locale",
                "de",  # German not supported
            ],
        )
        assert result.exit_code != 0

    def test_score_locale_short_option(self, sample_shigong_path, temp_dir):
        """Test score command with -l short option for locale."""
        docx_path = Path(temp_dir) / "report.docx"
        result = runner.invoke(
            app,
            [
                "score",
                "-i",
                sample_shigong_path,
                "--mode",
                "rules",
                "--docx-out",
                str(docx_path),
                "-l",
                "en",
            ],
        )
        assert result.exit_code == 0
        assert "DOCX report generated" in result.stdout

    def test_batch_locale_short_option(self, sample_shigong_path, temp_dir):
        """Test batch command with -l short option for locale."""
        result = runner.invoke(
            app,
            [
                "batch",
                "-i",
                sample_shigong_path,
                "-o",
                temp_dir,
                "-l",
                "en",
            ],
        )
        assert result.exit_code == 0
        assert "files to process" in result.stdout or "Found" in result.stdout
