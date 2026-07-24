"""Tests for web_ui module."""

from __future__ import annotations

import io
import os
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from fastapi.testclient import TestClient

try:
    import pymupdf
except Exception:  # noqa: BLE001 - environment-dependent optional dependency
    pymupdf = None

TEST_API_KEY = "test-auth-key-do-not-use"
AUTH_HEADERS = {"X-API-Key": TEST_API_KEY}
_ORIGINAL_API_KEYS = os.environ.get("API_KEYS")
os.environ["API_KEYS"] = TEST_API_KEY
try:
    from app.main import app
finally:
    if _ORIGINAL_API_KEYS is None:
        os.environ.pop("API_KEYS", None)
    else:
        os.environ["API_KEYS"] = _ORIGINAL_API_KEYS


@pytest.fixture(autouse=True)
def isolate_api_keys():
    """Keep UI integration tests independent from local authentication state."""
    with patch.dict(os.environ, {"API_KEYS": TEST_API_KEY}, clear=False):
        yield


class TestReadUploadedFile:
    """Tests for read_uploaded_file function."""

    def test_read_txt_file(self):
        """Test reading a .txt file."""
        from app.web_ui import read_uploaded_file

        mock_file = MagicMock()
        mock_file.name = "test.txt"
        mock_file.read.return_value = "测试文本内容".encode("utf-8")

        result = read_uploaded_file(mock_file)
        assert result == "测试文本内容"

    def test_read_docx_file(self):
        """Test reading a .docx file."""
        from app.web_ui import read_uploaded_file

        # Create a real DOCX in memory
        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()

        mock_file = MagicMock()
        mock_file.name = "test.docx"
        mock_file.read.return_value = docx_bytes

        result = read_uploaded_file(mock_file)
        assert "第一段内容" in result
        assert "第二段内容" in result

    def test_read_pdf_file(self):
        """Test reading a .pdf file."""
        from app.web_ui import read_uploaded_file

        if pymupdf is None:
            pytest.skip("PyMuPDF not available in this environment")

        # Create a real PDF in memory
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "PDF Test Content", fontname="helv")
        pdf_buffer = io.BytesIO()
        doc.save(pdf_buffer)
        doc.close()
        pdf_bytes = pdf_buffer.getvalue()

        mock_file = MagicMock()
        mock_file.name = "test.pdf"
        mock_file.read.return_value = pdf_bytes

        result = read_uploaded_file(mock_file)
        assert "PDF Test Content" in result

    def test_read_unsupported_format(self):
        """Test reading an unsupported file format raises error."""
        from app.web_ui import read_uploaded_file

        mock_file = MagicMock()
        mock_file.name = "test.xyz"

        with pytest.raises(ValueError, match="不支持的文件格式"):
            read_uploaded_file(mock_file)

    def test_read_uppercase_extension(self):
        """Test reading files with uppercase extensions."""
        from app.web_ui import read_uploaded_file

        mock_file = MagicMock()
        mock_file.name = "TEST.TXT"
        mock_file.read.return_value = "Upper case test".encode("utf-8")

        result = read_uploaded_file(mock_file)
        assert result == "Upper case test"


class TestWebUIModuleImport:
    """Tests for web_ui module structure."""

    def test_main_function_exists(self):
        """Test that main function exists."""
        from app.web_ui import main

        assert callable(main)

    def test_read_uploaded_file_exists(self):
        """Test that read_uploaded_file function exists."""
        from app.web_ui import read_uploaded_file

        assert callable(read_uploaded_file)


class TestWebUIIntegration:
    """Integration tests for web UI scoring flow."""

    def test_score_flow_with_txt_content(self):
        """Test the scoring flow with text content."""
        from app.config import load_config
        from app.engine.scorer import score_text

        config = load_config()
        text = "施工组织设计：本工程采用安全文明施工措施。"
        report = score_text(text, config.rubric, config.lexicon)

        assert hasattr(report, "total_score")
        assert hasattr(report, "dimension_scores")

    def test_export_report_integration(self, tmp_path: Path):
        """Test report export to DOCX."""
        from app.config import load_config
        from app.engine.docx_exporter import export_report_to_docx
        from app.engine.scorer import score_text

        config = load_config()
        text = "施工组织设计测试文档"
        report = score_text(text, config.rubric, config.lexicon)
        report_dict = report.model_dump()

        output_path = export_report_to_docx(report_dict, tmp_path / "report.docx")
        assert Path(output_path).exists()
        assert Path(output_path).stat().st_size > 0


class TestEmbeddedAuthControls:
    """Authentication controls and public redirect metadata boundaries."""

    def test_root_exposes_password_save_and_clear_controls(self):
        client = TestClient(app)
        response = client.get("/")

        assert response.status_code == 200
        page = response.text
        assert 'id="apiKeyInput" type="password"' in page
        assert 'id="saveApiKey"' in page
        assert 'id="clearApiKey"' in page
        assert 'id="apiKeyStatus"' in page
        assert 'localStorage.setItem("api_key", key)' in page
        assert 'localStorage.removeItem("api_key")' in page
        assert "if (input) input.value = '';" in page
        assert "setApiKeyStatus('已保存', false)" in page
        assert "未保存 key" in page
        assert "AUTH_KEY_MISSING" in page
        assert "AUTH_KEY_INVALID" in page
        assert "AUTH_NOT_CONFIGURED" in page

    def test_business_fetches_use_header_without_key_urls_or_console_output(self):
        client = TestClient(app)
        page = client.get("/").text

        assert "X-API-Key" in page
        assert "?api_key" not in page
        assert "console.log" not in page
        assert "__qingtianDownloadProtected" in page
        assert "await response.blob()" in page
        assert "await res.text()" in page

    @patch("app.main.create_project")
    def test_create_project_redirect_is_generic(self, create_project):
        create_project.return_value = SimpleNamespace(id="secret-project-id")
        client = TestClient(app)

        response = client.post(
            "/web/create_project",
            data={"name": "secret-project-name"},
            headers=AUTH_HEADERS,
            follow_redirects=False,
        )

        assert response.status_code == 303
        assert response.headers["location"] == "/?created=1"
        assert "secret-project-name" not in response.headers["location"]
        assert "secret-project-id" not in response.headers["location"]

    def test_root_ignores_project_metadata_query_values(self):
        client = TestClient(app)
        response = client.get(
            "/?created=1&create_ok=secret-project-name&project_id=secret-project-id"
        )

        assert response.status_code == 200
        assert "项目已创建，请使用 API key 刷新项目列表。" in response.text
        assert "secret-project-name" not in response.text
        assert "secret-project-id" not in response.text
