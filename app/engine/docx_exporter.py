"""DOCX 导出模块：将结构化评分报告渲染为 Word 文档。"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except Exception:
    Document = None
    WD_PARAGRAPH_ALIGNMENT = None

HIGH_PRIORITY = ["07", "09", "02", "03"]


def _safe_snippet(evidence: Dict[str, Any]) -> str:
    return (evidence or {}).get("snippet", "") or "未检索到证据片段"


def _truncate_cn(text: str, max_len: int = 80) -> str:
    if text is None:
        return ""
    text = text.strip()
    if len(text) <= max_len:
        return text
    return text[:max_len] + "…"


def _qingtian_comment(code: str, message: str) -> str:
    if code == "P-ACTION-001":
        return "措施缺'参数/频次/验收/责任'要素，落地性不足。"
    if code == "P-EMPTY-001":
        return "表述偏承诺型，缺可核查指标与闭环动作。"
    return (message or "")[:40]


def _top_penalties(report: Dict[str, Any], limit: int = 10) -> List[Dict[str, Any]]:
    penalties = report.get("penalties", []) or []
    penalties_sorted = sorted(penalties, key=lambda x: float(x.get("deduct", 0.0)), reverse=True)
    return penalties_sorted[:limit]


def export_report_to_docx(report: Dict[str, Any], output_path: str | Path) -> Path:
    """将评分报告导出为 DOCX 文件。

    Args:
        report: 评分报告 JSON 字典
        output_path: 输出文件路径

    Returns:
        输出文件的 Path 对象
    """
    if Document is None or WD_PARAGRAPH_ALIGNMENT is None:
        raise RuntimeError("DOCX 导出不可用：请安装与当前系统架构兼容的 python-docx/lxml。")
    doc = Document()
    output_path = Path(output_path)

    # 标题
    title = doc.add_heading("青天评标官终版报告", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 提取基本信息
    judge_mode = report.get("judge_mode", "unknown")
    judge_source = report.get("judge_source", "unknown")
    spark_called = report.get("spark_called", False)
    confidence = report.get("overall", {}).get("confidence_0_1")

    total_score = report.get("total_score")
    if total_score is None:
        total_score = report.get("overall", {}).get("total_score_0_100", 0)

    # A. 评分结论
    doc.add_heading("A. 评分结论", level=1)
    table_a = doc.add_table(rows=5, cols=2)
    table_a.style = "Table Grid"
    rows_data = [
        ("总分（0-100）", str(total_score)),
        ("置信度（0-1）", str(confidence) if confidence else "N/A"),
        ("judge_mode", str(judge_mode)),
        ("judge_source", str(judge_source)),
        ("spark_called", str(spark_called)),
    ]
    for i, (label, value) in enumerate(rows_data):
        table_a.rows[i].cells[0].text = label
        table_a.rows[i].cells[1].text = value
    doc.add_paragraph()

    # B. 高优维度诊断
    doc.add_heading("B. 高优维度诊断", level=1)
    doc.add_paragraph("高优维度权重策略：07/09/02/03")

    dim_scores = report.get("dimension_scores", {})
    for dim_id in HIGH_PRIORITY:
        dim = dim_scores.get(dim_id, {})
        dim_name = dim.get("name", f"维度{dim_id}")
        score = dim.get("score", 0)
        max_score = dim.get("max_score", 10)

        doc.add_heading(f"{dim_id} {dim_name}（{score}/{max_score}）", level=2)

        # 定义要点
        if spark_called:
            definition_list = dim.get("definition_points", [])[:3]
        else:
            hits = dim.get("hits", [])[:3]
            definition_list = hits if hits else ["未在文本中提取到明确要点"]
        doc.add_paragraph(f"定义要点：{'；'.join(definition_list) if definition_list else '无'}")

        # 缺陷
        if spark_called:
            defects_list = dim.get("defects", [])[:3]
        else:
            defects_list = ["参数/频次/验收/责任等落地要素表述不足"]
        doc.add_paragraph(f"缺陷：{'；'.join(defects_list)}")

        # 改进
        if spark_called:
            improvements_list = dim.get("improvements", [])[:3]
        else:
            improvements_list = ["建议补充可量化参数与验收闭环，并明确责任岗位与频次"]
        doc.add_paragraph(f"改进建议：{'；'.join(improvements_list)}")

        # 证据
        evidence_list = dim.get("evidence", [])[:2]
        evidence_text = "；".join(
            [_truncate_cn(_safe_snippet(e), 60) for e in (evidence_list or [{}])]
        )
        doc.add_paragraph(f"证据：{evidence_text if evidence_text else '无'}")

    # C. 扣分清单
    doc.add_heading("C. 扣分清单（Top 10）", level=1)
    top_penalties = _top_penalties(report, limit=10)

    if top_penalties:
        table_c = doc.add_table(rows=len(top_penalties) + 1, cols=4)
        table_c.style = "Table Grid"
        # 表头
        headers = ["编码", "扣分", "原因", "证据摘要"]
        for j, h in enumerate(headers):
            table_c.rows[0].cells[j].text = h
        # 数据行
        for i, p in enumerate(top_penalties, start=1):
            code = p.get("code", "")
            deduct = p.get("deduct", 0)
            message = p.get("message", "")
            evidence = _safe_snippet(p.get("evidence_span") or p.get("evidence", {}))
            evidence = _truncate_cn(evidence, 40)

            table_c.rows[i].cells[0].text = code
            table_c.rows[i].cells[1].text = str(deduct)
            table_c.rows[i].cells[2].text = message[:30] if message else ""
            table_c.rows[i].cells[3].text = evidence
    else:
        doc.add_paragraph("无扣分项。")
    doc.add_paragraph()

    # D. 改进建议
    doc.add_heading("D. 一次性提升清单", level=1)
    suggestions = report.get("suggestions", [])[:8]
    if suggestions:
        for s in suggestions:
            dim = s.get("dimension", "")
            action = s.get("action", "")
            gain = s.get("expected_gain", 0)
            doc.add_paragraph(f"• [{dim}] {action}（预计+{gain}分）")
    else:
        doc.add_paragraph("暂无改进建议。")

    # E. 附录
    doc.add_heading("E. 附录：证据索引说明", level=1)
    doc.add_paragraph("• start_index/end_index/snippet 表示原文区间与片段引用。")
    doc.add_paragraph(
        "• 如配置 OpenAI 凭证与 gpt-5.4，即可获得「青天 ChatGPT Judge 真评」并替换高优维度诊断内容。"
    )

    # 保存
    doc.save(str(output_path))
    return output_path
