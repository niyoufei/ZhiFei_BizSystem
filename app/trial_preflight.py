from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Dict, List, Mapping, Optional, Sequence

try:
    from docx import Document
except Exception:
    Document = None


def _as_dict(value: object) -> Dict[str, object]:
    return value if isinstance(value, dict) else {}


def _as_list(value: object) -> List[object]:
    return value if isinstance(value, list) else []


def _to_int(value: object) -> int:
    try:
        if value is None or value == "":
            return 0
        return int(float(value))
    except (TypeError, ValueError):
        return 0


def _to_float(value: object) -> Optional[float]:
    try:
        if value is None or value == "":
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _append_unique(target: List[str], text: object) -> None:
    value = str(text or "").strip()
    if value and value not in target:
        target.append(value)


def _system_improvement_closure_gate_label(gate_id: object, minimum_ready_projects: int) -> str:
    raw = str(gate_id or "").strip()
    if raw == "minimum_evaluated_projects":
        return f"具备真实样本的项目数达到 {minimum_ready_projects} 个"
    if raw == "minimum_ready_projects":
        return f"达到第一阶段就绪的项目数达到 {minimum_ready_projects} 个"
    if raw == "all_evaluated_projects_phase1_ready":
        return "所有具备真实样本的项目均已达到第一阶段就绪"
    if raw == "all_evaluated_projects_current_display_match_qt":
        return "所有具备真实样本的项目当前分均已对齐青天结果"
    if raw == "all_evaluated_projects_calibrator_stable":
        return "所有具备真实样本的项目当前项目级校准器均稳定"
    return raw or "-"


def _build_system_improvement_closure_gate_details(
    *,
    failed_gates: List[str],
    minimum_ready_projects: int,
    evaluated_project_count: int,
    ready_project_count: int,
    not_ready_project_count: int,
    current_display_matches_qt_pass_count: int,
    next_step_entrypoint_key: str,
    next_step_entrypoint_label: str,
    next_step_action_label: str,
    next_step_detail: str,
    next_priority_project_id: str,
    next_priority_project_name: str,
    next_candidate_project_id: str,
    next_candidate_project_name: str,
) -> List[Dict[str, object]]:
    gate_details: List[Dict[str, object]] = []
    default_entrypoint_key = next_step_entrypoint_key or "evaluation_summary"
    default_entrypoint_label = (
        next_step_entrypoint_label or "前往「5) 自我学习与进化」执行“跨项目汇总评估”"
    )
    default_action_label = next_step_action_label or "执行跨项目汇总评估"
    default_project_id = next_priority_project_id or next_candidate_project_id
    default_project_name = next_priority_project_name or next_candidate_project_name
    for gate_id in failed_gates:
        entrypoint_key = default_entrypoint_key
        entrypoint_label = default_entrypoint_label
        action_label = default_action_label
        project_id = default_project_id or None
        project_name = default_project_name or None
        summary = "系统总封关当前仍未满足该门。"
        detail = next_step_detail or "请按当前建议入口继续收口系统总封关前置条件。"
        raw_gate_id = str(gate_id or "").strip()
        if raw_gate_id == "minimum_evaluated_projects":
            summary = (
                f"当前仅有 {evaluated_project_count} 个项目进入跨项目评估，"
                f"至少需要 {minimum_ready_projects} 个。"
            )
            detail = next_step_detail or "需先推进候选项目进入真实评标与跨项目评估范围。"
        elif raw_gate_id == "minimum_ready_projects":
            summary = (
                f"当前仅有 {ready_project_count} 个项目达到第一阶段就绪，"
                f"至少需要 {minimum_ready_projects} 个。"
            )
            detail = next_step_detail or "需继续推进候选项目或优先收口未就绪项目。"
        elif raw_gate_id == "all_evaluated_projects_phase1_ready":
            summary = f"当前仍有 {not_ready_project_count} 个已评估项目未达到第一阶段就绪。"
            detail = next_step_detail or "建议优先收口当前最优先未就绪项目。"
        elif raw_gate_id == "all_evaluated_projects_current_display_match_qt":
            mismatch_count = max(0, evaluated_project_count - current_display_matches_qt_pass_count)
            summary = f"当前仍有 {mismatch_count} 个已评估项目的当前展示分未完全对齐青天结果。"
            detail = "建议先执行跨项目汇总评估定位偏差项目，再逐项收口当前分、校准器与闭环。"
            entrypoint_key = "evaluation_summary"
            entrypoint_label = "前往「5) 自我学习与进化」执行“跨项目汇总评估”"
            action_label = "执行跨项目汇总评估"
        elif raw_gate_id == "all_evaluated_projects_calibrator_stable":
            summary = "当前仍有已评估项目的项目级校准器未稳定。"
            detail = "建议先执行跨项目汇总评估并复核项目级校准器与一键闭环状态。"
            entrypoint_key = "evaluation_summary"
            entrypoint_label = "前往「5) 自我学习与进化」执行“跨项目汇总评估”"
            action_label = "执行跨项目汇总评估"
        gate_details.append(
            {
                "id": raw_gate_id,
                "label": _system_improvement_closure_gate_label(
                    raw_gate_id, minimum_ready_projects
                ),
                "summary": summary,
                "detail": detail,
                "project_id": project_id,
                "project_name": project_name,
                "entrypoint_key": entrypoint_key,
                "entrypoint_label": entrypoint_label,
                "action_label": action_label,
            }
        )
    return gate_details


def _build_system_improvement_project_gap_details(
    closure_summary: Mapping[str, object],
) -> List[Dict[str, object]]:
    gap_details: List[Dict[str, object]] = []
    not_ready_project_summaries = [
        _as_dict(item)
        for item in _as_list(closure_summary.get("not_ready_project_summaries"))
        if _as_dict(item)
    ]
    evaluated_project_summaries = [
        _as_dict(item)
        for item in _as_list(closure_summary.get("evaluated_project_summaries"))
        if _as_dict(item)
    ]
    seen: set[tuple[str, str]] = set()

    for row in not_ready_project_summaries:
        project_id = str(row.get("project_id") or "").strip()
        project_name = str(row.get("project_name") or "").strip() or project_id
        if not project_id:
            continue
        marker = ("phase1_ready_gap", project_id)
        if marker in seen:
            continue
        seen.add(marker)
        failed_gates = [
            str(item or "").strip()
            for item in _as_list(row.get("failed_gates"))
            if str(item or "").strip()
        ]
        detail = ""
        if failed_gates:
            detail = "未通过门：" + "、".join(failed_gates)
        recommendation = str(row.get("recommendation") or "").strip()
        if recommendation:
            detail = (detail + "；" if detail else "") + recommendation
        gap_details.append(
            {
                "id": f"phase1_ready_gap:{project_id}",
                "kind": "phase1_ready_gap",
                "kind_label": "第一阶段就绪缺口",
                "project_id": project_id,
                "project_name": project_name,
                "summary": (
                    f"当前项目仍未达到第一阶段就绪（未通过门 {int(row.get('failed_gate_count') or 0)} 个）。"
                ),
                "detail": detail or "建议优先收口该项目未通过门。",
                "entrypoint_key": str(row.get("entrypoint_key") or "").strip(),
                "entrypoint_label": str(row.get("entrypoint_label") or "").strip(),
                "action_label": str(row.get("action_label") or "").strip(),
            }
        )

    for row in evaluated_project_summaries:
        project_id = str(row.get("project_id") or "").strip()
        project_name = str(row.get("project_name") or "").strip() or project_id
        if not project_id or bool(row.get("current_display_matches_qt")):
            continue
        marker = ("current_display_matches_qt_gap", project_id)
        if marker in seen:
            continue
        seen.add(marker)
        detail_parts: List[str] = ["当前展示分尚未完全对齐青天结果。"]
        if row.get("current_mae_rmse_not_worse_than_v2") is False:
            detail_parts.append("当前 MAE/RMSE 仍未达到不劣于 V2。")
        if row.get("current_rank_corr_not_worse_than_v2") is False:
            detail_parts.append("当前排序相关性仍未达到不劣于 V2。")
        entrypoint_detail = str(row.get("entrypoint_detail") or "").strip()
        if entrypoint_detail:
            detail_parts.append(entrypoint_detail)
        gap_details.append(
            {
                "id": f"current_display_matches_qt_gap:{project_id}",
                "kind": "current_display_matches_qt_gap",
                "kind_label": "当前分对齐青天缺口",
                "project_id": project_id,
                "project_name": project_name,
                "summary": "当前项目的当前展示分仍未完全对齐青天结果。",
                "detail": " ".join(detail_parts).strip(),
                "entrypoint_key": str(row.get("entrypoint_key") or "").strip(),
                "entrypoint_label": str(row.get("entrypoint_label") or "").strip(),
                "action_label": str(row.get("action_label") or "").strip(),
            }
        )

    return gap_details


def _build_system_improvement_project_gate_gap_details(
    closure_summary: Mapping[str, object],
) -> List[Dict[str, object]]:
    gate_gap_details: List[Dict[str, object]] = []
    not_ready_project_summaries = [
        _as_dict(item)
        for item in _as_list(closure_summary.get("not_ready_project_summaries"))
        if _as_dict(item)
    ]
    seen: set[tuple[str, str]] = set()

    for row in not_ready_project_summaries:
        project_id = str(row.get("project_id") or "").strip()
        project_name = str(row.get("project_name") or "").strip() or project_id
        if not project_id:
            continue
        failed_gate_details = [
            _as_dict(item) for item in _as_list(row.get("failed_gate_details")) if _as_dict(item)
        ]
        if not failed_gate_details:
            failed_gate_details = [
                {"id": str(item or "").strip(), "label": str(item or "").strip()}
                for item in _as_list(row.get("failed_gates"))
                if str(item or "").strip()
            ]
        for gate in failed_gate_details:
            gate_id = str(gate.get("id") or "").strip()
            if not gate_id:
                continue
            marker = (project_id, gate_id)
            if marker in seen:
                continue
            seen.add(marker)
            gate_label = str(gate.get("label") or gate_id).strip() or gate_id
            gate_detail = str(gate.get("detail") or "").strip()
            entrypoint_key = str(
                gate.get("entrypoint_key") or row.get("entrypoint_key") or ""
            ).strip()
            entrypoint_label = str(
                gate.get("entrypoint_label") or row.get("entrypoint_label") or ""
            ).strip()
            action_label = str(gate.get("action_label") or row.get("action_label") or "").strip()
            detail_parts: List[str] = []
            if gate_detail:
                detail_parts.append(f"门内详情：{gate_detail}")
            recommendation = str(row.get("recommendation") or "").strip()
            if recommendation:
                detail_parts.append(recommendation)
            entrypoint_detail = str(
                gate.get("entrypoint_detail") or row.get("entrypoint_detail") or ""
            ).strip()
            if entrypoint_detail:
                detail_parts.append(entrypoint_detail)
            gate_gap_details.append(
                {
                    "id": f"{project_id}:{gate_id}",
                    "kind": "project_failed_gate",
                    "kind_label": "项目内未通过门",
                    "project_id": project_id,
                    "project_name": project_name,
                    "gate_id": gate_id,
                    "gate_label": gate_label,
                    "summary": f"当前项目仍未通过：{gate_label}。",
                    "detail": " ".join(detail_parts).strip() or "建议优先收口该门对应缺口。",
                    "entrypoint_key": entrypoint_key,
                    "entrypoint_label": entrypoint_label,
                    "action_label": action_label,
                }
            )

    return gate_gap_details


def _build_system_improvement_project_action_gap_details(
    closure_summary: Mapping[str, object],
) -> List[Dict[str, object]]:
    action_gap_details: List[Dict[str, object]] = []
    not_ready_project_summaries = [
        _as_dict(item)
        for item in _as_list(closure_summary.get("not_ready_project_summaries"))
        if _as_dict(item)
    ]

    for row in not_ready_project_summaries:
        project_id = str(row.get("project_id") or "").strip()
        project_name = str(row.get("project_name") or "").strip() or project_id
        if not project_id:
            continue
        failed_gate_details = [
            _as_dict(item) for item in _as_list(row.get("failed_gate_details")) if _as_dict(item)
        ]
        grouped: Dict[tuple[str, str, str], Dict[str, object]] = {}
        for gate in failed_gate_details:
            entrypoint_key = str(gate.get("entrypoint_key") or "").strip()
            entrypoint_label = str(gate.get("entrypoint_label") or "").strip()
            action_label = str(gate.get("action_label") or "").strip()
            if not entrypoint_key or not entrypoint_label:
                continue
            marker = (entrypoint_key, entrypoint_label, action_label)
            bucket = grouped.setdefault(
                marker,
                {
                    "entrypoint_key": entrypoint_key,
                    "entrypoint_label": entrypoint_label,
                    "action_label": action_label,
                    "gate_ids": [],
                    "gate_labels": [],
                    "entrypoint_details": [],
                },
            )
            gate_id = str(gate.get("id") or "").strip()
            gate_label = str(gate.get("label") or gate_id).strip() or gate_id
            if gate_id and gate_id not in bucket["gate_ids"]:
                bucket["gate_ids"].append(gate_id)
            if gate_label and gate_label not in bucket["gate_labels"]:
                bucket["gate_labels"].append(gate_label)
            entrypoint_detail = str(gate.get("entrypoint_detail") or "").strip()
            if entrypoint_detail and entrypoint_detail not in bucket["entrypoint_details"]:
                bucket["entrypoint_details"].append(entrypoint_detail)

        for entrypoint_key, entrypoint_label, action_label in sorted(grouped.keys()):
            bucket = grouped[(entrypoint_key, entrypoint_label, action_label)]
            gate_labels = [
                str(item) for item in bucket.get("gate_labels") or [] if str(item).strip()
            ]
            entrypoint_details = [
                str(item) for item in bucket.get("entrypoint_details") or [] if str(item).strip()
            ]
            gate_count = len(gate_labels)
            detail_parts: List[str] = []
            if gate_labels:
                detail_parts.append("关联未通过门：" + "、".join(gate_labels))
            if entrypoint_details:
                detail_parts.append("；".join(entrypoint_details))
            action_gap_details.append(
                {
                    "id": f"{project_id}:{entrypoint_key}",
                    "kind": "project_action_gap",
                    "kind_label": "按动作归并的项目收口",
                    "project_id": project_id,
                    "project_name": project_name,
                    "entrypoint_key": entrypoint_key,
                    "entrypoint_label": entrypoint_label,
                    "action_label": action_label,
                    "gate_count": gate_count,
                    "summary": (
                        f"当前项目有 {gate_count} 个未通过门建议通过该动作收口。"
                        if gate_count > 0
                        else "当前项目建议通过该动作继续收口。"
                    ),
                    "detail": " ".join(detail_parts).strip() or "建议优先执行该动作继续收口。",
                }
            )

    return action_gap_details


def _system_improvement_entrypoint_execution_mode(entrypoint_key: object) -> Dict[str, str]:
    raw = str(entrypoint_key or "").strip()
    if raw in {"auto_run_reflection"}:
        return {"execution_mode": "auto", "execution_mode_label": "自动闭环优先"}
    if raw in {
        "system_self_check",
        "system_data_hygiene",
        "evaluation_summary",
        "project_evaluation",
    }:
        return {"execution_mode": "readonly", "execution_mode_label": "只读诊断"}
    if raw in {"feedback_governance"}:
        return {"execution_mode": "review", "execution_mode_label": "人工复核优先"}
    return {"execution_mode": "manual", "execution_mode_label": "人工处理优先"}


def _system_improvement_action_group_meta(execution_mode: object) -> Dict[str, str]:
    raw = str(execution_mode or "").strip()
    if raw == "auto":
        return {
            "action_group": "auto",
            "action_group_label": "可自动收口动作",
            "group_reason_label": "该动作支持直接执行自动闭环收口。",
        }
    if raw == "readonly":
        return {
            "action_group": "readonly",
            "action_group_label": "只读诊断动作",
            "group_reason_label": "该动作仅用于只读诊断，不直接改写项目状态。",
        }
    return {
        "action_group": "manual",
        "action_group_label": "必须人工处理动作",
        "group_reason_label": "该动作需要人工录入、判断或复核后才能继续收口。",
    }


def _build_system_improvement_global_action_gap_details(
    project_action_gap_details: Sequence[Mapping[str, object]],
    *,
    next_priority_project_id: str = "",
    next_candidate_project_id: str = "",
) -> List[Dict[str, object]]:
    global_action_gap_details: List[Dict[str, object]] = []
    grouped: Dict[tuple[str, str, str], Dict[str, object]] = {}

    for item in project_action_gap_details:
        row = _as_dict(item)
        entrypoint_key = str(row.get("entrypoint_key") or "").strip()
        entrypoint_label = str(row.get("entrypoint_label") or "").strip()
        action_label = str(row.get("action_label") or "").strip()
        project_id = str(row.get("project_id") or "").strip()
        project_name = str(row.get("project_name") or "").strip() or project_id
        if not entrypoint_key or not entrypoint_label or not project_id:
            continue
        marker = (entrypoint_key, entrypoint_label, action_label)
        bucket = grouped.setdefault(
            marker,
            {
                "entrypoint_key": entrypoint_key,
                "entrypoint_label": entrypoint_label,
                "action_label": action_label,
                "project_ids": [],
                "project_names": [],
                "rows": [],
                "gate_count_total": 0,
            },
        )
        if project_id not in bucket["project_ids"]:
            bucket["project_ids"].append(project_id)
        if project_name and project_name not in bucket["project_names"]:
            bucket["project_names"].append(project_name)
        bucket["rows"].append(row)
        bucket["gate_count_total"] += max(0, _to_int(row.get("gate_count")))

    for entrypoint_key, entrypoint_label, action_label in sorted(grouped.keys()):
        bucket = grouped[(entrypoint_key, entrypoint_label, action_label)]
        project_ids = [str(item) for item in bucket.get("project_ids") or [] if str(item).strip()]
        project_names = [
            str(item) for item in bucket.get("project_names") or [] if str(item).strip()
        ]
        rows = [_as_dict(item) for item in bucket.get("rows") or [] if _as_dict(item)]
        gate_count_total = max(0, _to_int(bucket.get("gate_count_total")))
        priority_project_id = ""
        priority_project_name = ""
        priority_reason_label = ""
        priority_sort_label = ""
        for preferred_id in (next_priority_project_id, next_candidate_project_id):
            preferred_id = str(preferred_id or "").strip()
            if preferred_id and preferred_id in project_ids:
                priority_project_id = preferred_id
                if preferred_id == str(next_priority_project_id or "").strip():
                    priority_reason_label = "该项目是当前系统总封关优先收口项目。"
                    priority_sort_label = "系统总封关优先项目。"
                else:
                    priority_reason_label = "该项目是当前系统总封关候选首选项目。"
                    priority_sort_label = "系统总封关候选首选项目。"
                break
        if not priority_project_id and rows:
            sorted_rows = sorted(
                rows,
                key=lambda item: (
                    -_to_int(item.get("gate_count")),
                    str(item.get("project_name") or item.get("project_id") or ""),
                ),
            )
            winner = sorted_rows[0]
            priority_project_id = str(winner.get("project_id") or "").strip()
            priority_reason_label = f"该项目在此动作下关联未通过门最多（{max(0, _to_int(winner.get('gate_count')))} 个）。"
            priority_sort_label = "关联未通过门数量最多。"
        if not priority_project_id and project_ids:
            priority_project_id = project_ids[0]
        for row in rows:
            row_project_id = str(row.get("project_id") or "").strip()
            if row_project_id == priority_project_id:
                priority_project_name = (
                    str(row.get("project_name") or "").strip() or priority_project_id
                )
                break
        impacted_project_count = len(project_ids)
        if not priority_project_name and project_names:
            priority_project_name = project_names[0]
        if not priority_reason_label and impacted_project_count == 1:
            priority_reason_label = "该动作当前只影响这一个项目。"
            priority_sort_label = "当前仅影响单项目。"
        execution_mode_meta = _system_improvement_entrypoint_execution_mode(entrypoint_key)
        action_group_meta = _system_improvement_action_group_meta(
            execution_mode_meta["execution_mode"]
        )
        detail_parts: List[str] = []
        if project_names:
            detail_parts.append("涉及项目：" + "、".join(project_names))
        if priority_project_name:
            detail_parts.append(f"建议优先从“{priority_project_name}”开始。")
        row_details = []
        for row in rows:
            detail = str(row.get("detail") or "").strip()
            if detail and detail not in row_details:
                row_details.append(detail)
        if row_details:
            detail_parts.append("；".join(row_details[:3]))
        global_action_gap_details.append(
            {
                "id": f"global_action:{entrypoint_key}",
                "kind": "global_action_gap",
                "kind_label": "系统级优先动作",
                "project_id": priority_project_id,
                "project_name": priority_project_name,
                "entrypoint_key": entrypoint_key,
                "entrypoint_label": entrypoint_label,
                "action_label": action_label,
                "priority_reason_label": priority_reason_label,
                "execution_mode": execution_mode_meta["execution_mode"],
                "execution_mode_label": execution_mode_meta["execution_mode_label"],
                "action_group": action_group_meta["action_group"],
                "action_group_label": action_group_meta["action_group_label"],
                "group_reason_label": action_group_meta["group_reason_label"],
                "project_count": impacted_project_count,
                "gate_count_total": gate_count_total,
                "priority_sort_label": priority_sort_label,
                "summary": (
                    f"当前有 {impacted_project_count} 个项目、{gate_count_total} 个未通过门建议通过该动作收口。"
                    if gate_count_total > 0
                    else f"当前有 {impacted_project_count} 个项目建议通过该动作继续收口。"
                ),
                "detail": " ".join(detail_parts).strip() or "建议优先执行该系统级动作继续收口。",
            }
        )

    global_action_gap_details.sort(
        key=lambda item: (
            -_to_int(item.get("project_count")),
            -_to_int(item.get("gate_count_total")),
            str(item.get("action_label") or ""),
        )
    )
    return global_action_gap_details


def _build_system_improvement_focus_workstream_diagnostic_action_gap_details(
    focus_workstreams: Sequence[Mapping[str, object]],
    *,
    existing_global_action_gap_details: Sequence[Mapping[str, object]],
    next_priority_project_id: str = "",
    next_priority_project_name: str = "",
) -> List[Dict[str, object]]:
    existing_entrypoints = {
        str(_as_dict(item).get("entrypoint_key") or "").strip()
        for item in existing_global_action_gap_details
        if str(_as_dict(item).get("entrypoint_key") or "").strip()
    }
    grouped: Dict[tuple[str, str, str], Dict[str, object]] = {}
    for item in focus_workstreams:
        row = _as_dict(item)
        row_status = str(row.get("status") or "").strip()
        entrypoint_key = str(row.get("entrypoint_key") or "").strip()
        entrypoint_label = str(row.get("entrypoint_label") or "").strip()
        action_label = str(row.get("action_label") or "").strip()
        if row_status == "ok" or not entrypoint_key or not entrypoint_label:
            continue
        execution_mode_meta = _system_improvement_entrypoint_execution_mode(entrypoint_key)
        if execution_mode_meta["execution_mode"] != "readonly":
            continue
        if entrypoint_key in existing_entrypoints:
            continue
        marker = (entrypoint_key, entrypoint_label, action_label)
        bucket = grouped.setdefault(
            marker,
            {
                "entrypoint_key": entrypoint_key,
                "entrypoint_label": entrypoint_label,
                "action_label": action_label,
                "project_ids": [],
                "project_names": [],
                "titles": [],
                "summaries": [],
                "details": [],
            },
        )
        project_id = str(row.get("project_id") or "").strip()
        project_name = str(row.get("project_name") or "").strip() or project_id
        if project_id and project_id not in bucket["project_ids"]:
            bucket["project_ids"].append(project_id)
        if project_name and project_name not in bucket["project_names"]:
            bucket["project_names"].append(project_name)
        title = str(row.get("title") or "").strip()
        summary = str(row.get("summary") or "").strip()
        detail = str(row.get("detail") or "").strip()
        if title and title not in bucket["titles"]:
            bucket["titles"].append(title)
        if summary and summary not in bucket["summaries"]:
            bucket["summaries"].append(summary)
        if detail and detail not in bucket["details"]:
            bucket["details"].append(detail)

    diagnostic_rows: List[Dict[str, object]] = []
    for entrypoint_key, entrypoint_label, action_label in sorted(grouped.keys()):
        bucket = grouped[(entrypoint_key, entrypoint_label, action_label)]
        project_ids = [str(item) for item in bucket.get("project_ids") or [] if str(item).strip()]
        project_names = [
            str(item) for item in bucket.get("project_names") or [] if str(item).strip()
        ]
        titles = [str(item) for item in bucket.get("titles") or [] if str(item).strip()]
        summaries = [str(item) for item in bucket.get("summaries") or [] if str(item).strip()]
        details = [str(item) for item in bucket.get("details") or [] if str(item).strip()]
        priority_project_id = ""
        priority_project_name = ""
        priority_reason_label = ""
        priority_sort_label = ""
        if next_priority_project_id:
            priority_project_id = next_priority_project_id
            priority_project_name = next_priority_project_name or next_priority_project_id
            priority_reason_label = (
                "该动作用于继续诊断当前系统级缺口，建议先从当前优先收口项目开始。"
            )
            priority_sort_label = "系统总封关优先项目。"
        elif project_ids:
            priority_project_id = project_ids[0]
            priority_project_name = project_names[0] if project_names else priority_project_id
            priority_reason_label = "该动作用于继续诊断当前系统级缺口。"
            priority_sort_label = "当前诊断动作命中的首个项目。"
        execution_mode_meta = _system_improvement_entrypoint_execution_mode(entrypoint_key)
        action_group_meta = _system_improvement_action_group_meta(
            execution_mode_meta["execution_mode"]
        )
        summary = (
            f"当前有 {len(titles)} 条系统级诊断工作流建议先执行该动作。"
            if titles
            else "当前建议先执行该系统级诊断动作。"
        )
        detail_parts: List[str] = []
        if titles:
            detail_parts.append("关联工作流：" + "、".join(titles))
        if summaries:
            detail_parts.append("；".join(summaries[:2]))
        if details:
            detail_parts.append("；".join(details[:2]))
        if priority_project_name:
            detail_parts.append(f"建议优先从“{priority_project_name}”开始。")
        diagnostic_rows.append(
            {
                "id": f"global_action:{entrypoint_key}",
                "kind": "global_action_gap",
                "kind_label": "系统级优先动作",
                "project_id": priority_project_id,
                "project_name": priority_project_name,
                "entrypoint_key": entrypoint_key,
                "entrypoint_label": entrypoint_label,
                "action_label": action_label,
                "priority_reason_label": priority_reason_label,
                "execution_mode": execution_mode_meta["execution_mode"],
                "execution_mode_label": execution_mode_meta["execution_mode_label"],
                "action_group": action_group_meta["action_group"],
                "action_group_label": action_group_meta["action_group_label"],
                "group_reason_label": action_group_meta["group_reason_label"],
                "project_count": max(1, len(project_ids))
                if priority_project_id
                else len(project_ids),
                "gate_count_total": 0,
                "priority_sort_label": priority_sort_label,
                "summary": summary,
                "detail": " ".join(detail_parts).strip() or "建议先执行该系统级诊断动作。",
            }
        )
    return diagnostic_rows


def _split_system_improvement_global_action_gap_details(
    global_action_gap_details: Sequence[Mapping[str, object]],
) -> Dict[str, List[Dict[str, object]]]:
    auto_rows: List[Dict[str, object]] = []
    readonly_rows: List[Dict[str, object]] = []
    manual_rows: List[Dict[str, object]] = []
    for item in global_action_gap_details:
        row = _as_dict(item)
        if not row:
            continue
        action_group = str(row.get("action_group") or "").strip()
        if action_group == "auto":
            auto_rows.append(dict(row))
        elif action_group == "readonly":
            readonly_rows.append(dict(row))
        else:
            manual_rows.append(dict(row))
    return {
        "global_auto_action_gap_details": auto_rows,
        "global_readonly_action_gap_details": readonly_rows,
        "global_manual_action_gap_details": manual_rows,
    }


def _build_system_improvement_global_action_group_summaries(
    global_action_groups: Mapping[str, Sequence[Mapping[str, object]]],
) -> List[Dict[str, object]]:
    group_specs = [
        (
            "auto",
            "可自动收口动作",
            "global_auto_action_gap_details",
            "当前已有 {count} 条可自动收口动作，可优先通过自动闭环压缩系统缺口。",
            "当前未命中可自动收口动作，说明现有系统级缺口主要仍依赖人工处理或只读诊断。",
        ),
        (
            "readonly",
            "只读诊断动作",
            "global_readonly_action_gap_details",
            "当前已有 {count} 条只读诊断动作，适合先诊断再决定是否进入人工或自动收口。",
            "当前未命中只读诊断动作，说明现有系统级缺口已由自动或人工动作直接承接。",
        ),
        (
            "manual",
            "必须人工处理动作",
            "global_manual_action_gap_details",
            "当前已有 {count} 条必须人工处理动作，需人工录入、复核或治理后继续推进。",
            "当前未命中必须人工处理动作，说明现有系统级缺口暂不依赖人工录入、复核或治理。",
        ),
    ]
    summaries: List[Dict[str, object]] = []
    for action_group, action_group_label, key, active_template, empty_reason_label in group_specs:
        rows = [
            _as_dict(item) for item in _as_list(global_action_groups.get(key)) if _as_dict(item)
        ]
        count = len(rows)
        status = "active" if count > 0 else "empty"
        summaries.append(
            {
                "action_group": action_group,
                "action_group_label": action_group_label,
                "count": count,
                "status": status,
                "summary": (
                    active_template.format(count=count) if count > 0 else empty_reason_label
                ),
                "empty_reason_label": empty_reason_label if count <= 0 else "",
            }
        )
    return summaries


def _build_system_improvement_focus_workstream_status_summaries(
    focus_workstreams: Sequence[Mapping[str, object]],
) -> List[Dict[str, object]]:
    status_counts = {"ok": 0, "warn": 0, "blocked": 0}
    rows_by_status: Dict[str, List[Dict[str, object]]] = {"ok": [], "warn": [], "blocked": []}
    for item in focus_workstreams:
        row = _as_dict(item)
        row_status = str(row.get("status") or "").strip()
        if row_status in status_counts:
            status_counts[row_status] += 1
            rows_by_status[row_status].append(dict(row))

    def _priority_rank(row: Mapping[str, object], workstream_status: str) -> tuple[int, str]:
        workstream_id = str(row.get("id") or "").strip()
        if workstream_status == "blocked":
            rank_map = {
                "runtime_stability": 0,
                "data_hygiene": 1,
                "system_closure": 2,
                "priority_project": 3,
                "learning_alignment": 4,
            }
        elif workstream_status == "warn":
            rank_map = {
                "priority_project": 0,
                "system_closure": 1,
                "learning_alignment": 2,
                "runtime_stability": 3,
                "data_hygiene": 4,
            }
        else:
            rank_map = {
                "runtime_stability": 0,
                "data_hygiene": 1,
                "learning_alignment": 2,
                "system_closure": 3,
                "priority_project": 4,
            }
        return (rank_map.get(workstream_id, 99), workstream_id)

    status_specs = [
        (
            "ok",
            "正常",
            "当前已有 {count} 条工作流处于正常状态，可继续保持当前治理节奏。",
            "当前没有处于正常状态的工作流，说明系统级主工作流仍在持续收口。",
        ),
        (
            "warn",
            "待收口",
            "当前已有 {count} 条工作流仍待收口，建议继续按建议入口推进。",
            "当前没有待收口工作流，说明现有主工作流不是正常就是阻断。",
        ),
        (
            "blocked",
            "阻断",
            "当前已有 {count} 条工作流处于阻断状态，需优先排除阻断后再继续推进。",
            "当前没有阻断工作流，说明现有系统级缺口暂不阻断继续试车。",
        ),
    ]
    summaries: List[Dict[str, object]] = []
    for (
        workstream_status,
        workstream_status_label,
        active_template,
        empty_reason_label,
    ) in status_specs:
        count = status_counts[workstream_status]
        is_active = count > 0
        priority_row = (
            min(
                rows_by_status[workstream_status],
                key=lambda row: _priority_rank(row, workstream_status),
            )
            if rows_by_status[workstream_status]
            else {}
        )
        summaries.append(
            {
                "workstream_status": workstream_status,
                "workstream_status_label": workstream_status_label,
                "count": count,
                "status": "active" if is_active else "empty",
                "summary": (
                    active_template.format(count=count) if is_active else empty_reason_label
                ),
                "empty_reason_label": "" if is_active else empty_reason_label,
                "priority_workstream_id": str(priority_row.get("id") or "").strip(),
                "priority_workstream_title": str(priority_row.get("title") or "").strip(),
                "priority_project_id": str(priority_row.get("project_id") or "").strip(),
                "priority_project_name": str(priority_row.get("project_name") or "").strip(),
                "priority_entrypoint_key": str(priority_row.get("entrypoint_key") or "").strip(),
                "priority_entrypoint_label": str(
                    priority_row.get("entrypoint_label") or ""
                ).strip(),
                "priority_action_label": str(priority_row.get("action_label") or "").strip(),
            }
        )
    return summaries


def _ops_agent_quality_label(agent_name: str) -> str:
    raw = str(agent_name or "").strip()
    mapping = {
        "sre_watchdog": "SRE 看门狗",
        "data_hygiene": "数据卫生智能体",
        "runtime_repair": "运行态修复智能体",
        "project_flow": "项目流程智能体",
        "tender_project_flow": "招标流程智能体",
        "upload_flow": "上传流程智能体",
        "scoring_quality": "评分质量智能体",
        "evolution": "进化智能体",
        "learning_calibration": "学习校准智能体",
    }
    return mapping.get(raw, raw or "-")


def _build_system_improvement_ops_agent_quality_summary(
    ops_agents_status: Optional[Mapping[str, object]],
    ops_agents_history: Optional[Sequence[Mapping[str, object]]] = None,
) -> Dict[str, object]:
    snapshot = _as_dict(ops_agents_status)
    overall = _as_dict(snapshot.get("overall"))
    settings = _as_dict(snapshot.get("settings"))
    agents = {
        str(name or "").strip(): _as_dict(row)
        for name, row in _as_dict(snapshot.get("agents")).items()
        if str(name or "").strip() and _as_dict(row)
    }
    snapshot_path = str(snapshot.get("snapshot_path") or "").strip()
    load_error = str(snapshot.get("load_error") or "").strip()
    generated_at = str(snapshot.get("generated_at") or "").strip()
    snapshot_available = bool(
        generated_at and (agents or overall or _to_int(snapshot.get("agent_count")) > 0)
    )
    overall_status = str(overall.get("status") or "").strip()
    pass_count = _to_int(overall.get("pass_count"))
    warn_count = _to_int(overall.get("warn_count"))
    fail_count = _to_int(overall.get("fail_count"))
    agent_count = _to_int(snapshot.get("agent_count") or len(agents))
    duration_ms = _to_int(overall.get("duration_ms"))
    auto_repair_enabled = bool(settings.get("auto_repair"))
    auto_evolve_enabled = bool(settings.get("auto_evolve"))

    data_hygiene_agent = _as_dict(agents.get("data_hygiene"))
    runtime_repair_agent = _as_dict(agents.get("runtime_repair"))
    evolution_agent = _as_dict(agents.get("evolution"))
    learning_calibration_agent = _as_dict(agents.get("learning_calibration"))

    data_hygiene_actions = _as_dict(data_hygiene_agent.get("actions"))
    runtime_repair_actions = _as_dict(runtime_repair_agent.get("actions"))
    runtime_repair_metrics = _as_dict(runtime_repair_agent.get("metrics"))
    evolution_metrics = _as_dict(evolution_agent.get("metrics"))
    learning_metrics = _as_dict(learning_calibration_agent.get("metrics"))
    manual_confirmation_rows = [
        _as_dict(item)
        for item in _as_list(learning_calibration_agent.get("manual_confirmation_rows"))
        if _as_dict(item)
    ]

    auto_repair_attempted_count = 0
    auto_repair_success_count = 0
    for action_row in (
        _as_dict(data_hygiene_actions.get("repair")),
        _as_dict(runtime_repair_actions.get("repair_data_hygiene")),
        _as_dict(runtime_repair_actions.get("restart_runtime")),
    ):
        if not action_row:
            continue
        if bool(action_row.get("attempted")):
            auto_repair_attempted_count += 1
            if bool(action_row.get("ok")):
                auto_repair_success_count += 1

    auto_fixed_count = _to_int(runtime_repair_metrics.get("auto_fixed_count"))
    evolve_attempted_count = _to_int(learning_metrics.get("evolve_attempted_count"))
    evolve_success_count = _to_int(learning_metrics.get("evolve_success_count"))
    reflection_attempted_count = _to_int(learning_metrics.get("reflection_attempted_count"))
    reflection_success_count = _to_int(learning_metrics.get("reflection_success_count"))
    auto_evolve_attempted_count = evolve_attempted_count + reflection_attempted_count
    auto_evolve_success_count = evolve_success_count + reflection_success_count
    manual_confirmation_required_count = _to_int(
        learning_metrics.get("manual_confirmation_required_count")
    )
    post_verify_failed_count = _to_int(learning_metrics.get("post_verify_failed_count"))
    bootstrap_monitoring_count = _to_int(learning_metrics.get("bootstrap_monitoring_count"))
    llm_account_low_quality_pool_count = _to_int(
        learning_metrics.get("llm_account_low_quality_pool_count")
    )
    pending_evolve_after = _to_int(evolution_metrics.get("pending_evolve_after"))

    overall_status_label = "未知"
    if overall_status == "pass":
        overall_status_label = "巡检正常"
    elif overall_status == "warn":
        overall_status_label = "巡检待关注"
    elif overall_status == "fail":
        overall_status_label = "巡检失败"

    strengths: List[str] = []
    blockers: List[str] = []
    warnings: List[str] = []
    recommendations: List[str] = []

    if load_error:
        _append_unique(blockers, f"巡检快照读取失败：{load_error}。")
    elif not snapshot_available:
        _append_unique(warnings, "当前尚无自动巡检快照，无法判断最近一轮智能体工作质量。")
    else:
        _append_unique(
            strengths,
            f"最近一轮多智能体巡检已完成：通过 {pass_count} 个，待收口 {warn_count} 个，失败 {fail_count} 个。",
        )
        if fail_count <= 0:
            _append_unique(strengths, "最近一轮巡检未命中失败智能体。")
        else:
            _append_unique(
                blockers, f"最近一轮巡检存在 {fail_count} 个失败智能体，需先排除失败链路。"
            )

    if auto_repair_enabled:
        _append_unique(strengths, "自动修复开关已开启。")
    else:
        _append_unique(warnings, "当前自动修复开关未开启，巡检只能做只读诊断。")

    if auto_evolve_enabled:
        _append_unique(strengths, "自动学习/自动进化开关已开启。")
    else:
        _append_unique(warnings, "当前自动学习/自动进化开关未开启，自我优化链路不会自动推进。")

    if auto_fixed_count > 0:
        _append_unique(strengths, f"最近一轮已自动修复 {auto_fixed_count} 项运行态问题。")
    elif snapshot_available:
        _append_unique(strengths, "最近一轮未发现需要自动修复的运行态问题。")

    if auto_repair_attempted_count > 0 and auto_repair_success_count < auto_repair_attempted_count:
        _append_unique(
            warnings,
            f"最近一轮自动修复已尝试 {auto_repair_attempted_count} 次，但仅成功 {auto_repair_success_count} 次。",
        )

    if auto_evolve_attempted_count > 0:
        if auto_evolve_success_count == auto_evolve_attempted_count:
            _append_unique(
                strengths,
                f"最近一轮自动学习/闭环共执行 {auto_evolve_attempted_count} 次，全部成功。",
            )
        else:
            _append_unique(
                warnings,
                f"最近一轮自动学习/闭环共执行 {auto_evolve_attempted_count} 次，但仅成功 {auto_evolve_success_count} 次。",
            )
    elif pending_evolve_after > 0:
        _append_unique(
            warnings,
            f"最近一轮后仍有 {pending_evolve_after} 个项目待继续推进自动进化。",
        )

    if manual_confirmation_required_count > 0:
        _append_unique(
            warnings,
            f"当前仍有 {manual_confirmation_required_count} 个项目需人工确认后，自动学习链路才能继续推进。",
        )
        first_manual_confirmation_row = (
            manual_confirmation_rows[0] if manual_confirmation_rows else {}
        )
        first_manual_confirmation_project_name = str(
            first_manual_confirmation_row.get("project_name")
            or first_manual_confirmation_row.get("project_id")
            or ""
        ).strip()
        first_manual_confirmation_detail = str(
            first_manual_confirmation_row.get("detail") or ""
        ).strip()
        if first_manual_confirmation_project_name:
            _append_unique(
                warnings,
                f"当前人工确认主项目为“{first_manual_confirmation_project_name}”"
                + (
                    f"：{first_manual_confirmation_detail}。"
                    if first_manual_confirmation_detail
                    else "。"
                ),
            )
    if bootstrap_monitoring_count > 0:
        _append_unique(
            warnings,
            f"当前仍有 {bootstrap_monitoring_count} 个项目处于小样本自举监控期，自动优化仍需继续观察。",
        )
    if llm_account_low_quality_pool_count > 0:
        _append_unique(
            warnings,
            f"当前有 {llm_account_low_quality_pool_count} 个服务提供方账号池历史质量分偏低，系统虽会自动降优先级，但智能体质量仍需继续观察。",
        )
    if post_verify_failed_count > 0:
        _append_unique(
            blockers,
            f"最近一轮自动学习/反射后复验仍有 {post_verify_failed_count} 处失败，不能视为最优稳定状态。",
        )

    for text in _as_list(snapshot.get("recommendations")):
        msg = str(text or "").strip()
        if msg:
            _append_unique(recommendations, msg)

    agent_rows: List[Dict[str, object]] = []
    for agent_name, agent_row in agents.items():
        row_status = str(agent_row.get("status") or "").strip()
        row_recommendations = [
            str(item or "").strip()
            for item in _as_list(agent_row.get("recommendations"))
            if str(item or "").strip()
        ]
        if row_status == "pass" and not row_recommendations:
            continue
        agent_rows.append(
            {
                "name": agent_name,
                "label": _ops_agent_quality_label(agent_name),
                "status": row_status or "unknown",
                "recommendation": row_recommendations[0] if row_recommendations else "",
            }
        )
    agent_rows.sort(
        key=lambda item: (
            {"fail": 0, "warn": 1, "pass": 2}.get(str(item.get("status") or ""), 9),
            str(item.get("name") or ""),
        )
    )

    history_rows = [_as_dict(item) for item in _as_list(ops_agents_history) if _as_dict(item)]
    recent_audit_rows = history_rows[-5:]
    recent_cycle_count = len(history_rows)
    recent_pass_cycle_count = sum(
        1 for item in history_rows if str(item.get("overall_status") or "").strip() == "pass"
    )
    recent_warn_cycle_count = sum(
        1 for item in history_rows if str(item.get("overall_status") or "").strip() == "warn"
    )
    recent_fail_cycle_count = sum(
        1 for item in history_rows if str(item.get("overall_status") or "").strip() == "fail"
    )
    total_auto_repair_attempted_count = sum(
        _to_int(item.get("auto_repair_attempted_count")) for item in history_rows
    )
    total_auto_repair_success_count = sum(
        _to_int(item.get("auto_repair_success_count")) for item in history_rows
    )
    total_auto_evolve_attempted_count = sum(
        _to_int(item.get("auto_evolve_attempted_count")) for item in history_rows
    )
    total_auto_evolve_success_count = sum(
        _to_int(item.get("auto_evolve_success_count")) for item in history_rows
    )
    repair_success_rate = (
        total_auto_repair_success_count / total_auto_repair_attempted_count
        if total_auto_repair_attempted_count > 0
        else None
    )
    evolve_success_rate = (
        total_auto_evolve_success_count / total_auto_evolve_attempted_count
        if total_auto_evolve_attempted_count > 0
        else None
    )
    recent_non_pass_streak_count = 0
    for item in reversed(history_rows):
        if str(item.get("overall_status") or "").strip() == "pass":
            break
        recent_non_pass_streak_count += 1
    recent_manual_gate_cycle_count = sum(
        1 for item in history_rows if _to_int(item.get("manual_confirmation_required_count")) > 0
    )
    recent_post_verify_failed_cycle_count = sum(
        1 for item in history_rows if _to_int(item.get("post_verify_failed_count")) > 0
    )
    recent_quality_reason_counts: Dict[str, Dict[str, object]] = {}
    for item in history_rows:
        reason_code = str(item.get("quality_reason_code") or "").strip()
        reason_label = str(item.get("quality_reason_label") or "").strip()
        if not reason_code:
            continue
        bucket = recent_quality_reason_counts.setdefault(
            reason_code,
            {
                "quality_reason_code": reason_code,
                "quality_reason_label": reason_label or reason_code,
                "count": 0,
            },
        )
        bucket["count"] = _to_int(bucket.get("count")) + 1
    recent_quality_reason_summary_rows = sorted(
        recent_quality_reason_counts.values(),
        key=lambda item: (
            -_to_int(_as_dict(item).get("count")),
            str(_as_dict(item).get("quality_reason_code") or ""),
        ),
    )
    latest_quality_reason_label = (
        str(history_rows[-1].get("quality_reason_label") or "").strip() if history_rows else ""
    )
    latest_quality_reason_code = (
        str(history_rows[-1].get("quality_reason_code") or "").strip() if history_rows else ""
    )
    latest_quality_reason_project_id = (
        str(history_rows[-1].get("quality_reason_project_id") or "").strip() if history_rows else ""
    )
    latest_quality_reason_project_name = (
        str(history_rows[-1].get("quality_reason_project_name") or "").strip()
        if history_rows
        else ""
    )
    latest_quality_reason_project_detail = (
        str(history_rows[-1].get("quality_reason_project_detail") or "").strip()
        if history_rows
        else ""
    )
    if latest_quality_reason_code == "manual_confirmation_required" and manual_confirmation_rows:
        first_manual_confirmation_row = manual_confirmation_rows[0]
        latest_quality_reason_project_id = (
            latest_quality_reason_project_id
            or str(first_manual_confirmation_row.get("project_id") or "").strip()
        )
        latest_quality_reason_project_name = (
            latest_quality_reason_project_name
            or str(
                first_manual_confirmation_row.get("project_name")
                or latest_quality_reason_project_id
                or ""
            ).strip()
        )
        latest_quality_reason_project_detail = (
            latest_quality_reason_project_detail
            or str(first_manual_confirmation_row.get("detail") or "").strip()
        )
    recent_same_reason_streak_count = 0
    if latest_quality_reason_code:
        for item in reversed(history_rows):
            if str(item.get("quality_reason_code") or "").strip() != latest_quality_reason_code:
                break
            recent_same_reason_streak_count += 1

    if recent_cycle_count > 0:
        _append_unique(
            strengths,
            f"最近 {recent_cycle_count} 轮巡检中，通过 {recent_pass_cycle_count} 轮、待收口 {recent_warn_cycle_count} 轮、失败 {recent_fail_cycle_count} 轮。",
        )
    if recent_non_pass_streak_count > 1:
        _append_unique(
            warnings,
            f"最近已连续 {recent_non_pass_streak_count} 轮巡检未进入 pass，自动巡检质量仍在持续待收口。",
        )
    if recent_manual_gate_cycle_count > 0:
        _append_unique(
            warnings,
            f"最近 {recent_manual_gate_cycle_count} 轮巡检命中过人工确认门，自动学习仍未进入全自动稳定区间。",
        )
    if recent_post_verify_failed_cycle_count > 0:
        _append_unique(
            blockers,
            f"最近 {recent_post_verify_failed_cycle_count} 轮巡检命中过复验失败，自动优化仍存在未闭环风险。",
        )
    if recent_same_reason_streak_count > 1 and latest_quality_reason_label:
        _append_unique(
            warnings,
            f"最近已连续 {recent_same_reason_streak_count} 轮巡检主因都为“{latest_quality_reason_label}”，说明该类问题尚未真正收口。",
        )
    if repair_success_rate is not None:
        _append_unique(
            strengths if repair_success_rate >= 1.0 else warnings,
            f"最近巡检历史中的自动修复成功率为 {repair_success_rate:.0%}（{total_auto_repair_success_count}/{total_auto_repair_attempted_count}）。",
        )
    if evolve_success_rate is not None:
        _append_unique(
            strengths if evolve_success_rate >= 1.0 else warnings,
            f"最近巡检历史中的自动学习成功率为 {evolve_success_rate:.0%}（{total_auto_evolve_success_count}/{total_auto_evolve_attempted_count}）。",
        )

    quality_status = "ready"
    quality_status_label = "自动巡检质量正常"
    if blockers:
        quality_status = "blocked"
        quality_status_label = "自动巡检存在失败或复验未收口项"
    elif warnings:
        quality_status = "watch"
        quality_status_label = "自动巡检可继续运行，但质量仍需关注"

    summary_label = (
        f"{quality_status_label} / 通过 {pass_count} / 待收口 {warn_count} / 失败 {fail_count}"
    )
    if auto_repair_attempted_count > 0 or auto_fixed_count > 0:
        summary_label += f" / 自动修复 {auto_fixed_count}"
    if auto_evolve_attempted_count > 0:
        summary_label += (
            f" / 自动学习成功 {auto_evolve_success_count}/{auto_evolve_attempted_count}"
        )

    return {
        "snapshot_available": snapshot_available,
        "snapshot_path": snapshot_path or None,
        "generated_at": generated_at or None,
        "load_error": load_error or None,
        "overall_status": overall_status or "unknown",
        "overall_status_label": overall_status_label,
        "quality_status": quality_status,
        "quality_status_label": quality_status_label,
        "summary_label": summary_label,
        "agent_count": agent_count,
        "pass_count": pass_count,
        "warn_count": warn_count,
        "fail_count": fail_count,
        "duration_ms": duration_ms,
        "recent_audit_rows": recent_audit_rows,
        "recent_cycle_count": recent_cycle_count,
        "recent_pass_cycle_count": recent_pass_cycle_count,
        "recent_warn_cycle_count": recent_warn_cycle_count,
        "recent_fail_cycle_count": recent_fail_cycle_count,
        "recent_non_pass_streak_count": recent_non_pass_streak_count,
        "recent_manual_gate_cycle_count": recent_manual_gate_cycle_count,
        "recent_post_verify_failed_cycle_count": recent_post_verify_failed_cycle_count,
        "latest_quality_reason_code": latest_quality_reason_code or None,
        "latest_quality_reason_label": latest_quality_reason_label or None,
        "latest_quality_reason_project_id": latest_quality_reason_project_id or None,
        "latest_quality_reason_project_name": latest_quality_reason_project_name or None,
        "latest_quality_reason_project_detail": latest_quality_reason_project_detail or None,
        "recent_same_reason_streak_count": recent_same_reason_streak_count,
        "recent_quality_reason_summary_rows": recent_quality_reason_summary_rows,
        "auto_repair_enabled": auto_repair_enabled,
        "auto_evolve_enabled": auto_evolve_enabled,
        "auto_repair_attempted_count": auto_repair_attempted_count,
        "auto_repair_success_count": auto_repair_success_count,
        "auto_fixed_count": auto_fixed_count,
        "total_auto_repair_attempted_count": total_auto_repair_attempted_count,
        "total_auto_repair_success_count": total_auto_repair_success_count,
        "repair_success_rate": repair_success_rate,
        "auto_evolve_attempted_count": auto_evolve_attempted_count,
        "auto_evolve_success_count": auto_evolve_success_count,
        "total_auto_evolve_attempted_count": total_auto_evolve_attempted_count,
        "total_auto_evolve_success_count": total_auto_evolve_success_count,
        "evolve_success_rate": evolve_success_rate,
        "manual_confirmation_required_count": manual_confirmation_required_count,
        "manual_confirmation_rows": manual_confirmation_rows[:8],
        "post_verify_failed_count": post_verify_failed_count,
        "bootstrap_monitoring_count": bootstrap_monitoring_count,
        "llm_account_low_quality_pool_count": llm_account_low_quality_pool_count,
        "agent_rows": agent_rows[:8],
        "strengths": strengths,
        "blockers": blockers,
        "warnings": warnings,
        "recommendations": recommendations[:8],
    }


def _build_trial_preflight_signoff(
    *,
    trial_run_ready: bool,
    status: str,
    status_label: str,
    blockers: List[str],
    warnings: List[str],
    metrics: Mapping[str, object],
) -> Dict[str, object]:
    if status == "blocked" or not trial_run_ready:
        decision = "hold"
        decision_label = "暂缓试车"
        risk_level = "high"
        risk_label = "高风险"
    elif status == "watch":
        decision = "approve_with_watch"
        decision_label = "建议试车（带警告）"
        risk_level = "medium"
        risk_label = "中风险"
    else:
        decision = "approve"
        decision_label = "建议试车"
        risk_level = "low"
        risk_label = "低风险"

    verification_checklist = [
        {
            "name": "系统自检",
            "passed": bool(metrics.get("self_check_ok")),
            "detail": "运行时必需项全部正常"
            if bool(metrics.get("self_check_ok"))
            else "存在未通过必需项",
        },
        {
            "name": "项目评分前置",
            "passed": bool(metrics.get("project_ready_to_score")),
            "detail": (
                "当前项目已满足评分前置"
                if bool(metrics.get("project_ready_to_score"))
                else "当前项目尚未满足评分前置"
            ),
        },
        {
            "name": "项目资料门禁",
            "passed": bool(metrics.get("project_gate_passed")),
            "detail": "资料门禁通过"
            if bool(metrics.get("project_gate_passed"))
            else "资料门禁未通过",
        },
        {
            "name": "项目级学习闭环",
            "passed": int(metrics.get("matched_score_record_count") or 0) > 0,
            "detail": (
                f"已形成 {int(metrics.get('matched_score_record_count') or 0)} 条有效评分记录匹配"
                if int(metrics.get("matched_score_record_count") or 0) > 0
                else "尚未形成有效评分记录匹配"
            ),
        },
        {
            "name": "项目级校准器",
            "passed": not bool(metrics.get("current_calibrator_degraded")),
            "detail": (
                str(metrics.get("current_calibrator_version") or "未部署项目级校准器")
                + ("（已退化）" if bool(metrics.get("current_calibrator_degraded")) else "")
            ),
        },
        {
            "name": "系统总封关",
            "passed": bool(metrics.get("system_closure_ready")),
            "detail": (
                "系统总封关前置条件已满足"
                if bool(metrics.get("system_closure_ready"))
                else "系统总封关仍未完成，但不阻断当前项目试车"
            ),
        },
    ]

    summary_label = f"{decision_label} / {risk_label} / 阻断 {len(blockers)} / 警告 {len(warnings)} / 状态 {status_label}"

    return {
        "decision": decision,
        "decision_label": decision_label,
        "risk_level": risk_level,
        "risk_label": risk_label,
        "summary_label": summary_label,
        "verification_checklist": verification_checklist,
    }


def _build_trial_preflight_record_draft(
    *,
    generated_at: str,
    signoff: Mapping[str, object],
    warnings: List[str],
    recommendations: List[str],
) -> Dict[str, object]:
    signoff_label = str(signoff.get("decision_label") or "-").strip() or "-"
    risk_label = str(signoff.get("risk_label") or "-").strip() or "-"
    warning_ack_required = bool(warnings)
    confirmation_hint = (
        "当前存在警告项，试车后请人工补记确认意见并逐条复核警告项。"
        if warning_ack_required
        else "试车完成后请补记执行人与结论确认。"
    )
    summary_label = f"待人工确认 / {signoff_label} / {risk_label}"
    return {
        "status": "pending_manual_confirmation",
        "status_label": "待人工确认",
        "summary_label": summary_label,
        "suggested_executed_at": generated_at,
        "executor_name": "待填写",
        "recommended_conclusion": signoff_label,
        "recommended_risk_label": risk_label,
        "warning_ack_required": warning_ack_required,
        "warning_ack_items": list(warnings),
        "confirmation_hint": confirmation_hint,
        "next_recommended_action": str(recommendations[0] or "").strip() if recommendations else "",
    }


def _material_conflict_kind_label(value: object) -> str:
    raw = str(value or "").strip()
    if raw == "numeric_mismatch":
        return "数值不一致"
    if raw == "term_coverage_missing":
        return "术语覆盖缺失"
    if raw == "material_consistency_missing":
        return "资料一致性缺失"
    return raw or "-"


def _material_review_entrypoint(material_type: object, material_type_label: str) -> Dict[str, str]:
    raw = str(material_type or "").strip()
    if raw == "tender_qa":
        return {
            "material_review_entrypoint_label": "前往「3) 项目资料」核对招标文件和答疑",
            "material_review_entrypoint_anchor": "#uploadMaterial",
        }
    if raw == "boq":
        return {
            "material_review_entrypoint_label": "前往「3) 项目资料」核对清单",
            "material_review_entrypoint_anchor": "#uploadMaterialBoq",
        }
    if raw == "drawing":
        return {
            "material_review_entrypoint_label": "前往「3) 项目资料」核对图纸",
            "material_review_entrypoint_anchor": "#uploadMaterialDrawing",
        }
    if raw == "site_photo":
        return {
            "material_review_entrypoint_label": "前往「3) 项目资料」核对现场照片",
            "material_review_entrypoint_anchor": "#uploadMaterialPhoto",
        }
    return {
        "material_review_entrypoint_label": f"前往「3) 项目资料」核对{material_type_label}",
        "material_review_entrypoint_anchor": "#section-materials",
    }


def _build_trial_preflight_material_review_reason(
    *,
    material_type_label: str,
    conflict_kind: object,
) -> str:
    conflict_kind_raw = str(conflict_kind or "").strip()
    if conflict_kind_raw == "numeric_mismatch":
        return f"当前该高严重度冲突来自{material_type_label}数值不一致，建议先核对对应资料来源文件和量化约束。"
    if conflict_kind_raw == "term_coverage_missing":
        return f"当前该高严重度冲突来自{material_type_label}术语覆盖缺失，建议先核对对应资料来源文件和关键术语表述。"
    return f"当前该高严重度冲突与{material_type_label}资料一致性相关，建议先核对对应资料来源文件。"


def _build_trial_preflight_shigong_entrypoint(
    *,
    scoring_readiness: Mapping[str, object],
) -> Dict[str, str]:
    submissions = _as_dict(scoring_readiness.get("submissions"))
    non_empty = _to_int(submissions.get("non_empty"))
    scored = _to_int(submissions.get("scored"))
    ready = bool(scoring_readiness.get("ready"))
    gate_passed = bool(scoring_readiness.get("gate_passed"))
    if non_empty <= 0:
        return {
            "entrypoint_key": "upload_shigong",
            "entrypoint_label": "前往「4) 项目施组」上传施组",
            "entrypoint_anchor": "#section-shigong",
            "entrypoint_reason_label": "当前项目还没有可用施组，需先上传施组后才能进入评分。",
        }
    if ready and gate_passed and scored < non_empty:
        return {
            "entrypoint_key": "score_shigong",
            "entrypoint_label": "前往「4) 项目施组」评分施组",
            "entrypoint_anchor": "#section-shigong",
            "entrypoint_reason_label": "当前项目已有待评分施组，且评分前置与资料门禁已满足，可直接评分施组。",
        }
    return {
        "entrypoint_key": "upload_shigong",
        "entrypoint_label": "前往「4) 项目施组」上传新版施组",
        "entrypoint_anchor": "#section-shigong",
        "entrypoint_reason_label": "当前项目已有已评分施组；若已按冲突项修改内容，应先上传新版施组，再重新评分。",
    }


def _build_trial_preflight_conflict_action(
    *,
    material_type_label: str,
    conflict_kind: object,
    shigong_entrypoint: Mapping[str, object],
) -> Dict[str, str]:
    conflict_kind_raw = str(conflict_kind or "").strip()
    if conflict_kind_raw == "numeric_mismatch":
        action_label = f"优先回到施组，补齐与{material_type_label}一致的量化约束、工程量或参数。"
    elif conflict_kind_raw == "term_coverage_missing":
        action_label = f"优先回到施组，补齐与{material_type_label}一致的术语、章节标题或措施表述。"
    else:
        action_label = f"优先回到施组，补齐与{material_type_label}一致的关键约束。"
    return {
        "entrypoint_key": str(shigong_entrypoint.get("entrypoint_key") or "upload_shigong"),
        "entrypoint_label": str(
            shigong_entrypoint.get("entrypoint_label") or "前往「4) 项目施组」上传施组"
        ),
        "entrypoint_anchor": str(shigong_entrypoint.get("entrypoint_anchor") or "#section-shigong"),
        "entrypoint_reason_label": str(
            shigong_entrypoint.get("entrypoint_reason_label")
            or "请先按当前项目状态选择合适的施组处理入口。"
        ),
        "action_label": action_label,
        "secondary_hint": f"必要时再回看「3) 项目资料」中的{material_type_label}来源文件是否齐全。",
    }


def _build_trial_preflight_warning_details(
    *,
    scoring_diagnostic: Mapping[str, object],
    scoring_readiness: Mapping[str, object],
) -> Dict[str, object]:
    evidence_trace = _as_dict(scoring_diagnostic.get("evidence_trace"))
    material_conflicts = _as_dict(evidence_trace.get("material_conflicts"))
    shigong_entrypoint = _build_trial_preflight_shigong_entrypoint(
        scoring_readiness=scoring_readiness,
    )
    high_conflicts: List[Dict[str, object]] = []
    for item in _as_list(material_conflicts.get("conflicts")):
        row = _as_dict(item)
        if str(row.get("severity") or "").strip() != "high":
            continue
        dimension_id = str(row.get("dimension_id") or "").strip() or "-"
        material_type = str(row.get("material_type") or "").strip()
        material_type_label = str(row.get("material_type_label") or "").strip() or "-"
        conflict_kind_label = _material_conflict_kind_label(row.get("conflict_kind"))
        label = str(row.get("label") or "").strip() or "-"
        action_meta = _build_trial_preflight_conflict_action(
            material_type_label=material_type_label,
            conflict_kind=row.get("conflict_kind"),
            shigong_entrypoint=shigong_entrypoint,
        )
        review_meta = _material_review_entrypoint(material_type, material_type_label)
        material_review_reason_label = _build_trial_preflight_material_review_reason(
            material_type_label=material_type_label,
            conflict_kind=row.get("conflict_kind"),
        )
        detail_parts: List[str] = []
        if bool(row.get("mandatory")):
            detail_parts.append("强制项")
        term_need = _to_int(row.get("term_need"))
        if term_need > 0:
            detail_parts.append(f"术语 {_to_int(row.get('term_hit'))}/{term_need}")
        num_need = _to_int(row.get("num_need"))
        if num_need > 0:
            detail_parts.append(f"数值 {_to_int(row.get('num_hit'))}/{num_need}")
        source_mode = str(row.get("source_mode") or "").strip()
        if source_mode:
            detail_parts.append(f"来源 {source_mode}")
        high_conflicts.append(
            {
                "dimension_id": dimension_id,
                "material_type": material_type,
                "material_type_label": material_type_label,
                "conflict_kind_label": conflict_kind_label,
                "label": label,
                "summary_label": (
                    f"维度{dimension_id} / {material_type_label} / {conflict_kind_label} / {label}"
                ),
                "detail_label": "；".join(detail_parts),
                "entrypoint_key": action_meta["entrypoint_key"],
                "entrypoint_label": action_meta["entrypoint_label"],
                "entrypoint_anchor": action_meta["entrypoint_anchor"],
                "entrypoint_reason_label": action_meta["entrypoint_reason_label"],
                "action_label": action_meta["action_label"],
                "material_review_entrypoint_label": review_meta["material_review_entrypoint_label"],
                "material_review_entrypoint_anchor": review_meta[
                    "material_review_entrypoint_anchor"
                ],
                "material_review_reason_label": material_review_reason_label,
                "secondary_hint": action_meta["secondary_hint"],
            }
        )

    recommendations = [
        str(x or "").strip()
        for x in _as_list(material_conflicts.get("recommendations"))
        if str(x or "").strip()
    ]
    return {
        "high_severity_material_conflict_count": len(high_conflicts),
        "high_severity_material_conflicts": high_conflicts[:5],
        "material_conflict_recommendations": recommendations[:5],
    }


def build_trial_preflight_report(
    *,
    base_url: str,
    project_id: str,
    project_name: str,
    self_check: Mapping[str, object],
    scoring_readiness: Mapping[str, object],
    mece_audit: Mapping[str, object],
    evolution_health: Mapping[str, object],
    scoring_diagnostic: Mapping[str, object],
    evaluation_summary: Mapping[str, object],
    data_hygiene: Mapping[str, object],
) -> Dict[str, object]:
    self_summary = _as_dict(self_check.get("summary"))
    mece_overall = _as_dict(mece_audit.get("overall"))
    evolution_summary = _as_dict(evolution_health.get("summary"))
    evolution_drift = _as_dict(evolution_health.get("drift"))
    scoring_summary = _as_dict(scoring_diagnostic.get("summary"))
    closure_summary = _as_dict(evaluation_summary.get("total_closure_readiness"))

    blockers: List[str] = []
    warnings: List[str] = []
    strengths: List[str] = []
    recommendations: List[str] = []

    self_check_ok = bool(self_check.get("ok"))
    project_ready = bool(scoring_readiness.get("ready"))
    gate_passed = bool(scoring_readiness.get("gate_passed"))
    mece_level = str(mece_overall.get("level") or "").strip() or "unknown"
    mece_health_score = _to_float(mece_overall.get("health_score"))
    gt_count = _to_int(evolution_summary.get("ground_truth_count"))
    matched_score_record_count = _to_int(evolution_summary.get("matched_score_record_count"))
    matched_prediction_count = _to_int(evolution_summary.get("matched_prediction_count"))
    calibrator_version = str(evolution_summary.get("current_calibrator_version") or "").strip()
    calibrator_degraded = bool(evolution_summary.get("current_calibrator_degraded"))
    evolution_weights_usable = bool(evolution_summary.get("evolution_weights_usable"))
    drift_level = str(evolution_drift.get("level") or "").strip() or "unknown"
    latest_score_confidence_level = str(
        scoring_summary.get("latest_score_confidence_level") or ""
    ).strip()
    conflict_count = _to_int(scoring_summary.get("material_conflict_high_severity_count"))
    closure_ready = bool(closure_summary.get("ready"))
    closure_failed_gates = [
        str(item or "").strip()
        for item in _as_list(closure_summary.get("failed_gates"))
        if str(item or "").strip()
    ]
    orphan_records_total = _to_int(data_hygiene.get("orphan_records_total"))

    if self_check_ok:
        _append_unique(strengths, "系统自检通过，运行时必需项全部正常。")
    else:
        failed_required = [
            str(item or "").strip()
            for item in _as_list(self_summary.get("failed_required_items"))
            if str(item or "").strip()
        ]
        detail = "、".join(failed_required) if failed_required else "存在未通过的必需项"
        _append_unique(blockers, f"系统自检未通过：{detail}。")
        _append_unique(recommendations, "先执行一次 doctor，并修复自检失败项后再试车。")

    if orphan_records_total > 0:
        _append_unique(blockers, f"数据卫生未通过：仍存在 {orphan_records_total} 条孤儿记录。")
        _append_unique(recommendations, "先处理数据卫生孤儿记录，再继续试车。")
    else:
        _append_unique(strengths, "数据卫生良好，未发现跨项目孤儿记录。")

    if project_ready and gate_passed:
        _append_unique(strengths, "当前项目评分前置与资料门禁均已通过。")
    else:
        issue_list = [
            str(item or "").strip()
            for item in _as_list(scoring_readiness.get("issues"))
            if str(item or "").strip()
        ]
        detail = "；".join(issue_list) if issue_list else "评分前置或资料门禁未通过"
        _append_unique(blockers, f"当前项目尚未满足评分前置：{detail}。")
        _append_unique(recommendations, "先补齐当前项目资料解析与评分前置，再进行试车。")

    if mece_level == "critical":
        _append_unique(blockers, "当前项目 MECE 诊断为 critical，主链路仍存在阻断项。")
        for item in _as_list(mece_audit.get("recommendations")):
            _append_unique(recommendations, item)
    elif mece_level == "watch":
        score_text = "-" if mece_health_score is None else f"{mece_health_score:.1f}"
        _append_unique(warnings, f"当前项目 MECE 健康度为 watch（{score_text}）。")
        for item in _as_list(mece_audit.get("recommendations")):
            _append_unique(recommendations, item)
    elif mece_level == "good":
        score_text = "-" if mece_health_score is None else f"{mece_health_score:.1f}"
        _append_unique(strengths, f"当前项目 MECE 健康度良好（{score_text}）。")

    if latest_score_confidence_level:
        _append_unique(
            strengths,
            f"最新施组评分置信度为 {latest_score_confidence_level}，智能评分链路可用。",
        )
    if conflict_count > 0:
        _append_unique(
            warnings,
            f"最新施组仍存在 {conflict_count} 个高严重度资料冲突，建议按评分证据链逐条收口。",
        )
    for item in _as_list(scoring_diagnostic.get("recommendations")):
        _append_unique(recommendations, item)

    if gt_count <= 0:
        _append_unique(warnings, "当前项目尚无真实评标样本，自我学习闭环尚未启动。")
        _append_unique(recommendations, "试车通过后尽快录入真实评标样本，启动项目级学习闭环。")
    else:
        _append_unique(strengths, f"当前项目已录入 {gt_count} 条真实评标样本。")

    if gt_count > 0 and matched_score_record_count <= 0:
        _append_unique(blockers, "真实评标样本尚未形成有效评分记录匹配，自我学习链路不可用。")
        _append_unique(recommendations, "先修复真实评标与评分记录关联，再继续试车。")
    elif matched_score_record_count > 0:
        _append_unique(
            strengths,
            f"真实评标已形成 {matched_score_record_count} 条有效评分记录匹配。",
        )

    if gt_count > 0 and matched_prediction_count <= 0:
        _append_unique(warnings, "真实评标样本尚未形成有效预测匹配，漂移评估仍偏弱。")
    elif matched_prediction_count > 0:
        _append_unique(
            strengths,
            f"真实评标已形成 {matched_prediction_count} 条有效预测匹配。",
        )

    if calibrator_degraded:
        _append_unique(blockers, "当前项目级校准器已退化，建议先治理后再试车。")
        _append_unique(recommendations, "优先执行一键闭环或评分治理，恢复稳定校准器后再试车。")
    elif calibrator_version:
        _append_unique(strengths, f"当前项目级校准器已部署：{calibrator_version}。")

    if evolution_weights_usable:
        _append_unique(strengths, "演化权重已可用，自我进化链路处于启用状态。")
    elif gt_count > 0:
        _append_unique(warnings, "已有真实样本，但演化权重当前不可用。")
        _append_unique(recommendations, "复核进化权重生成与部署链路。")

    if drift_level == "low":
        _append_unique(strengths, "近期漂移等级为 low。")
    elif drift_level == "insufficient_data":
        _append_unique(warnings, "近期漂移判断仍为 insufficient_data，学习成熟度仍处于小样本阶段。")
        _append_unique(recommendations, "继续累计真实评标样本，提升漂移判断与自学习稳定性。")
    elif drift_level and drift_level != "unknown":
        _append_unique(warnings, f"近期漂移等级为 {drift_level}，建议持续观测。")

    if closure_ready:
        _append_unique(strengths, "系统总封关前置条件已满足。")
    else:
        _append_unique(
            warnings,
            "系统总封关前置条件尚未全部满足，但这不阻断当前项目试车。",
        )
        next_step_title = str(closure_summary.get("next_step_title") or "").strip()
        next_step_detail = str(closure_summary.get("next_step_detail") or "").strip()
        if next_step_title:
            _append_unique(recommendations, f"{next_step_title}：{next_step_detail or '-'}")
        if closure_failed_gates:
            _append_unique(
                warnings,
                "系统总封关未通过门：" + "、".join(closure_failed_gates) + "。",
            )

    trial_run_ready = len(blockers) == 0
    if blockers:
        status = "blocked"
        status_label = "暂不建议试车"
    elif warnings:
        status = "watch"
        status_label = "可试车，但建议先关注警告项"
    else:
        status = "ready"
        status_label = "可直接试车"

    metrics = {
        "self_check_ok": self_check_ok,
        "project_ready_to_score": project_ready,
        "project_gate_passed": gate_passed,
        "project_mece_level": mece_level,
        "project_mece_health_score": mece_health_score,
        "ground_truth_count": gt_count,
        "matched_prediction_count": matched_prediction_count,
        "matched_score_record_count": matched_score_record_count,
        "current_calibrator_version": calibrator_version or None,
        "current_calibrator_degraded": calibrator_degraded,
        "evolution_weights_usable": evolution_weights_usable,
        "drift_level": drift_level,
        "latest_score_confidence_level": latest_score_confidence_level or None,
        "material_conflict_high_severity_count": conflict_count,
        "system_closure_ready": closure_ready,
        "system_closure_failed_gates": closure_failed_gates,
        "orphan_records_total": orphan_records_total,
    }
    signoff = _build_trial_preflight_signoff(
        trial_run_ready=trial_run_ready,
        status=status,
        status_label=status_label,
        blockers=blockers,
        warnings=warnings,
        metrics=metrics,
    )
    warning_details = _build_trial_preflight_warning_details(
        scoring_diagnostic=scoring_diagnostic,
        scoring_readiness=scoring_readiness,
    )

    generated_at = datetime.now(timezone.utc).isoformat()
    record_draft = _build_trial_preflight_record_draft(
        generated_at=generated_at,
        signoff=signoff,
        warnings=warnings,
        recommendations=recommendations,
    )

    return {
        "generated_at": generated_at,
        "base_url": base_url,
        "project_id": project_id,
        "project_name": project_name,
        "trial_run_ready": trial_run_ready,
        "status": status,
        "status_label": status_label,
        "metrics": metrics,
        "signoff": signoff,
        "warning_details": warning_details,
        "record_draft": record_draft,
        "strengths": strengths,
        "blockers": blockers,
        "warnings": warnings,
        "recommendations": recommendations,
    }


def build_system_improvement_overview_report(
    *,
    base_url: str,
    self_check: Mapping[str, object],
    data_hygiene: Mapping[str, object],
    evaluation_summary: Mapping[str, object],
    ops_agents_status: Optional[Mapping[str, object]] = None,
    ops_agents_history: Optional[Sequence[Mapping[str, object]]] = None,
) -> Dict[str, object]:
    self_summary = _as_dict(self_check.get("summary"))
    closure_summary = _as_dict(evaluation_summary.get("total_closure_readiness"))
    acceptance_pass_count = _as_dict(evaluation_summary.get("acceptance_pass_count"))
    blockers: List[str] = []
    warnings: List[str] = []
    strengths: List[str] = []
    recommendations: List[str] = []

    self_check_ok = bool(self_check.get("ok"))
    failed_required_count = _to_int(self_check.get("failed_required_count"))
    failed_optional_count = _to_int(self_check.get("failed_optional_count"))
    orphan_records_total = _to_int(data_hygiene.get("orphan_records_total"))
    project_count = _to_int(evaluation_summary.get("project_count"))
    evaluated_project_count = _to_int(closure_summary.get("evaluated_project_count"))
    ready_project_count = _to_int(closure_summary.get("ready_project_count"))
    not_ready_project_count = _to_int(closure_summary.get("not_ready_project_count"))
    candidate_project_count = _to_int(closure_summary.get("candidate_project_count"))
    minimum_ready_projects = _to_int(closure_summary.get("minimum_ready_projects"))
    system_closure_ready = bool(closure_summary.get("ready"))
    failed_gates = [
        str(item or "").strip()
        for item in _as_list(closure_summary.get("failed_gates"))
        if str(item or "").strip()
    ]
    next_step_title = str(closure_summary.get("next_step_title") or "").strip()
    next_step_detail = str(closure_summary.get("next_step_detail") or "").strip()
    next_step_entrypoint_key = str(closure_summary.get("next_step_entrypoint_key") or "").strip()
    next_step_entrypoint_label = str(
        closure_summary.get("next_step_entrypoint_label") or ""
    ).strip()
    next_step_action_label = str(closure_summary.get("next_step_action_label") or "").strip()
    blocker_kind = str(closure_summary.get("blocker_kind") or "").strip()
    next_priority_project_id = str(closure_summary.get("next_priority_project_id") or "").strip()
    next_priority_project_name = str(
        closure_summary.get("next_priority_project_name") or ""
    ).strip()
    next_candidate_project_id = str(closure_summary.get("next_candidate_project_id") or "").strip()
    next_candidate_project_name = str(
        closure_summary.get("next_candidate_project_name") or ""
    ).strip()
    current_display_matches_qt_pass_count = _to_int(
        acceptance_pass_count.get("current_display_matches_qt")
    )
    current_mae_rmse_not_worse_pass_count = _to_int(
        acceptance_pass_count.get("current_mae_rmse_not_worse_than_v2")
    )
    current_rank_corr_not_worse_pass_count = _to_int(
        acceptance_pass_count.get("current_rank_corr_not_worse_vs_v2")
    )
    affected_dataset_count = sum(
        1
        for item in _as_list(data_hygiene.get("datasets"))
        if _to_int(_as_dict(item).get("orphan_count")) > 0
    )

    if project_count <= 0:
        _append_unique(blockers, "当前系统暂无项目，无法执行真实试车与跨项目闭环评估。")
        _append_unique(recommendations, "先创建并推进至少 1 个真实项目，再执行系统级体检。")
    else:
        _append_unique(strengths, f"当前系统已纳管 {project_count} 个项目。")

    if self_check_ok:
        _append_unique(strengths, "系统自检通过，核心运行项正常。")
    else:
        failed_required = [
            str(item or "").strip()
            for item in _as_list(self_summary.get("failed_required_items"))
            if str(item or "").strip()
        ]
        detail = "、".join(failed_required) if failed_required else "存在未通过的核心自检项"
        _append_unique(blockers, f"系统自检未通过：{detail}。")
        _append_unique(recommendations, "先执行一次系统自检并修复核心失败项，再继续推进系统试车。")

    if failed_optional_count > 0:
        _append_unique(
            warnings,
            f"系统仍有 {failed_optional_count} 项降级告警，建议继续收口运行时稳定性。",
        )
        _append_unique(recommendations, "优先处理系统自检中的降级告警，避免试车期放大运行波动。")
    else:
        _append_unique(strengths, "系统未发现额外降级告警。")

    if orphan_records_total > 0:
        _append_unique(
            blockers,
            f"数据卫生未通过：仍有 {orphan_records_total} 条孤儿记录，影响 {affected_dataset_count} 个数据集。",
        )
        _append_unique(recommendations, "先执行数据卫生修复，再继续推进跨项目评估与总封关。")
    else:
        _append_unique(strengths, "数据卫生良好，未发现跨项目孤儿记录。")

    if evaluated_project_count <= 0:
        _append_unique(warnings, "当前尚无进入跨项目评估的真实项目，系统级学习成熟度仍不足。")
        _append_unique(recommendations, "先让至少 1 个真实项目完成评分、真实评标录入与项目级闭环。")
    else:
        _append_unique(
            strengths,
            f"当前已有 {evaluated_project_count} 个项目进入跨项目评估范围。",
        )

    if ready_project_count > 0:
        _append_unique(strengths, f"当前已有 {ready_project_count} 个项目达到第一阶段就绪。")
    if not_ready_project_count > 0:
        detail = (
            f"；当前最优先收口项目为“{next_priority_project_name}”。"
            if next_priority_project_name
            else "。"
        )
        _append_unique(
            warnings,
            f"当前仍有 {not_ready_project_count} 个项目未达到第一阶段就绪{detail}",
        )
    if candidate_project_count > 0:
        detail = (
            f"；当前最优先推进候选项目为“{next_candidate_project_name}”。"
            if next_candidate_project_name
            else "。"
        )
        _append_unique(
            warnings,
            f"当前仍有 {candidate_project_count} 个候选项目待继续推进闭环{detail}",
        )

    if system_closure_ready:
        _append_unique(strengths, "系统总封关前置条件已满足。")
    else:
        _append_unique(
            warnings,
            "系统总封关前置条件尚未全部满足，仍需继续收口跨项目闭环。",
        )
        if failed_gates:
            _append_unique(warnings, "系统总封关未通过门：" + "、".join(failed_gates) + "。")
        if next_step_title:
            _append_unique(recommendations, f"{next_step_title}：{next_step_detail or '-'}")

    if evaluated_project_count > 0:
        if current_display_matches_qt_pass_count == evaluated_project_count:
            _append_unique(strengths, "所有已评估项目的当前展示分均已对齐青天结果。")
        else:
            mismatch_count = max(
                0,
                evaluated_project_count - current_display_matches_qt_pass_count,
            )
            _append_unique(
                warnings,
                f"仍有 {mismatch_count} 个已评估项目的当前展示分未完全对齐青天结果。",
            )
            _append_unique(recommendations, "优先处理当前展示分与青天结果仍有偏差的项目。")

        if current_mae_rmse_not_worse_pass_count == evaluated_project_count:
            _append_unique(strengths, "所有已评估项目当前 MAE/RMSE 均未劣于 V2 基线。")
        else:
            lagging_count = max(
                0,
                evaluated_project_count - current_mae_rmse_not_worse_pass_count,
            )
            _append_unique(
                warnings,
                f"仍有 {lagging_count} 个已评估项目当前 MAE/RMSE 未达到不劣于 V2 的目标。",
            )
            _append_unique(recommendations, "优先对误差仍偏高的项目执行一键闭环与校准治理。")

        if current_rank_corr_not_worse_pass_count == evaluated_project_count:
            _append_unique(strengths, "所有已评估项目当前排序相关性均未劣于 V2 基线。")
        else:
            lagging_rank_count = max(
                0,
                evaluated_project_count - current_rank_corr_not_worse_pass_count,
            )
            _append_unique(
                warnings,
                f"仍有 {lagging_rank_count} 个已评估项目当前排序相关性未达到不劣于 V2 的目标。",
            )
            _append_unique(recommendations, "优先复核排序仍偏弱项目的当前分、校准器与规则闭环。")

    status = "ready"
    status_label = "系统级体检通过"
    if blockers:
        status = "blocked"
        status_label = "系统仍有阻断项，暂不建议进入系统级试车"
    elif warnings:
        status = "watch"
        status_label = "系统可继续试车，但仍需按清单持续收口"

    overall_ready = status == "ready"
    summary_label = (
        f"{status_label} / 阻断 {len(blockers)} / 警告 {len(warnings)} / 项目 {project_count}"
    )
    closure_project_id = ""
    closure_project_name = ""
    if blocker_kind == "close_not_ready_project" and next_priority_project_id:
        closure_project_id = next_priority_project_id
        closure_project_name = next_priority_project_name
    elif blocker_kind == "advance_candidate_project" and next_candidate_project_id:
        closure_project_id = next_candidate_project_id
        closure_project_name = next_candidate_project_name

    focus_workstreams: List[Dict[str, object]] = [
        {
            "id": "runtime_stability",
            "title": "运行稳定性",
            "status": "blocked"
            if not self_check_ok
            else ("warn" if failed_optional_count > 0 else "ok"),
            "summary": (
                "系统自检未通过，需先修复核心运行项。"
                if not self_check_ok
                else (
                    f"系统自检已通过，但仍有 {failed_optional_count} 项降级告警。"
                    if failed_optional_count > 0
                    else "系统自检已通过，运行稳定性处于正常状态。"
                )
            ),
            "detail": (
                f"核心失败 {failed_required_count} 项；降级告警 {failed_optional_count} 项。"
            ),
            "entrypoint_key": "system_self_check",
            "entrypoint_label": "前往「5) 自我学习与进化」执行“系统自检”",
            "action_label": "执行系统自检",
        },
        {
            "id": "data_hygiene",
            "title": "数据卫生",
            "status": "blocked" if orphan_records_total > 0 else "ok",
            "summary": (
                f"仍有 {orphan_records_total} 条孤儿记录待修复。"
                if orphan_records_total > 0
                else "当前未发现孤儿记录。"
            ),
            "detail": (
                f"影响数据集 {affected_dataset_count} 个。"
                if orphan_records_total > 0
                else "当前跨项目数据集状态干净。"
            ),
            "entrypoint_key": "system_data_hygiene",
            "entrypoint_label": "前往「5) 自我学习与进化」执行“数据卫生巡检”",
            "action_label": "执行数据卫生巡检",
        },
        {
            "id": "learning_alignment",
            "title": "跨项目学习对齐",
            "status": (
                "warn"
                if evaluated_project_count <= 0
                or current_display_matches_qt_pass_count < evaluated_project_count
                or current_mae_rmse_not_worse_pass_count < evaluated_project_count
                or current_rank_corr_not_worse_pass_count < evaluated_project_count
                else "ok"
            ),
            "summary": (
                "当前尚无进入跨项目评估的真实项目。"
                if evaluated_project_count <= 0
                else (
                    f"当前分对齐 {current_display_matches_qt_pass_count}/{evaluated_project_count}；"
                    f"误差不劣于 V2 {current_mae_rmse_not_worse_pass_count}/{evaluated_project_count}；"
                    f"排序不劣于 V2 {current_rank_corr_not_worse_pass_count}/{evaluated_project_count}。"
                )
            ),
            "detail": ("需继续通过真实评标样本、项目级闭环和校准器治理提升跨项目稳定性。"),
            "entrypoint_key": "evaluation_summary",
            "entrypoint_label": "前往「5) 自我学习与进化」执行“跨项目汇总评估”",
            "action_label": "查看跨项目汇总",
        },
        {
            "id": "system_closure",
            "title": "系统总封关",
            "status": "ok" if system_closure_ready else "warn",
            "summary": (
                "系统总封关前置条件已满足。"
                if system_closure_ready
                else (
                    str(closure_summary.get("status_label") or "").strip() or "系统总封关仍未完成。"
                )
            ),
            "detail": next_step_detail
            or (
                "当前可进入系统总封关核查。"
                if system_closure_ready
                else "仍需继续推进系统总封关前置条件。"
            ),
            "project_id": closure_project_id or None,
            "project_name": closure_project_name or None,
            "entrypoint_key": next_step_entrypoint_key or "evaluation_summary",
            "entrypoint_label": next_step_entrypoint_label
            or "前往「5) 自我学习与进化」执行“跨项目汇总评估”",
            "action_label": next_step_action_label or "查看跨项目汇总",
        },
    ]

    priority_project_id = next_priority_project_id or next_candidate_project_id
    priority_project_name = next_priority_project_name or next_candidate_project_name
    if priority_project_name:
        focus_workstreams.append(
            {
                "id": "priority_project",
                "title": "当前优先收口对象",
                "status": "warn",
                "summary": (
                    f"优先收口项目“{next_priority_project_name}”。"
                    if next_priority_project_name
                    else f"优先推进候选项目“{next_candidate_project_name}”。"
                ),
                "detail": next_step_detail or "按系统总封关建议优先推进当前最关键项目。",
                "project_id": priority_project_id or None,
                "project_name": priority_project_name or None,
                "entrypoint_key": next_step_entrypoint_key or "",
                "entrypoint_label": next_step_entrypoint_label or "",
                "action_label": next_step_action_label or "",
            }
        )

    closure_gate_details = _build_system_improvement_closure_gate_details(
        failed_gates=failed_gates,
        minimum_ready_projects=minimum_ready_projects,
        evaluated_project_count=evaluated_project_count,
        ready_project_count=ready_project_count,
        not_ready_project_count=not_ready_project_count,
        current_display_matches_qt_pass_count=current_display_matches_qt_pass_count,
        next_step_entrypoint_key=next_step_entrypoint_key,
        next_step_entrypoint_label=next_step_entrypoint_label,
        next_step_action_label=next_step_action_label,
        next_step_detail=next_step_detail,
        next_priority_project_id=next_priority_project_id,
        next_priority_project_name=next_priority_project_name,
        next_candidate_project_id=next_candidate_project_id,
        next_candidate_project_name=next_candidate_project_name,
    )
    project_gap_details = _build_system_improvement_project_gap_details(closure_summary)
    project_gate_gap_details = _build_system_improvement_project_gate_gap_details(closure_summary)
    project_action_gap_details = _build_system_improvement_project_action_gap_details(
        closure_summary
    )
    global_action_gap_details = _build_system_improvement_global_action_gap_details(
        project_action_gap_details,
        next_priority_project_id=next_priority_project_id,
        next_candidate_project_id=next_candidate_project_id,
    )
    global_action_gap_details.extend(
        _build_system_improvement_focus_workstream_diagnostic_action_gap_details(
            focus_workstreams,
            existing_global_action_gap_details=global_action_gap_details,
            next_priority_project_id=next_priority_project_id,
            next_priority_project_name=next_priority_project_name,
        )
    )
    global_action_gap_details.sort(
        key=lambda item: (
            -_to_int(_as_dict(item).get("project_count")),
            -_to_int(_as_dict(item).get("gate_count_total")),
            str(_as_dict(item).get("action_label") or ""),
        )
    )
    global_action_groups = _split_system_improvement_global_action_gap_details(
        global_action_gap_details
    )
    global_action_group_summaries = _build_system_improvement_global_action_group_summaries(
        global_action_groups
    )
    focus_workstream_status_summaries = _build_system_improvement_focus_workstream_status_summaries(
        focus_workstreams
    )
    ops_agent_quality_summary = _build_system_improvement_ops_agent_quality_summary(
        ops_agents_status,
        ops_agents_history=ops_agents_history,
    )

    metrics = {
        "self_check_ok": self_check_ok,
        "failed_required_count": failed_required_count,
        "failed_optional_count": failed_optional_count,
        "orphan_records_total": orphan_records_total,
        "affected_dataset_count": affected_dataset_count,
        "project_count": project_count,
        "evaluated_project_count": evaluated_project_count,
        "ready_project_count": ready_project_count,
        "not_ready_project_count": not_ready_project_count,
        "candidate_project_count": candidate_project_count,
        "minimum_ready_projects": minimum_ready_projects,
        "system_closure_ready": system_closure_ready,
        "system_closure_failed_gates": failed_gates,
        "current_display_matches_qt_pass_count": current_display_matches_qt_pass_count,
        "current_mae_rmse_not_worse_pass_count": current_mae_rmse_not_worse_pass_count,
        "current_rank_corr_not_worse_pass_count": current_rank_corr_not_worse_pass_count,
        "next_priority_project_id": next_priority_project_id or None,
        "next_priority_project_name": next_priority_project_name or None,
        "next_candidate_project_id": next_candidate_project_id or None,
        "next_candidate_project_name": next_candidate_project_name or None,
    }

    return {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "base_url": base_url,
        "overall_ready": overall_ready,
        "status": status,
        "status_label": status_label,
        "summary_label": summary_label,
        "metrics": metrics,
        "focus_workstreams": focus_workstreams,
        "focus_workstream_status_summaries": focus_workstream_status_summaries,
        "ops_agent_quality_summary": ops_agent_quality_summary,
        "closure_gate_details": closure_gate_details,
        "project_gap_details": project_gap_details,
        "project_gate_gap_details": project_gate_gap_details,
        "project_action_gap_details": project_action_gap_details,
        "global_action_gap_details": global_action_gap_details,
        "global_action_group_summaries": global_action_group_summaries,
        "global_auto_action_gap_details": global_action_groups["global_auto_action_gap_details"],
        "global_readonly_action_gap_details": global_action_groups[
            "global_readonly_action_gap_details"
        ],
        "global_manual_action_gap_details": global_action_groups[
            "global_manual_action_gap_details"
        ],
        "strengths": strengths,
        "blockers": blockers,
        "warnings": warnings,
        "recommendations": recommendations,
    }


def render_trial_preflight_markdown(report: Mapping[str, object]) -> str:
    metrics = _as_dict(report.get("metrics"))
    signoff = _as_dict(report.get("signoff"))
    warning_details = _as_dict(report.get("warning_details"))
    record_draft = _as_dict(report.get("record_draft"))
    strengths = [
        str(x or "").strip() for x in _as_list(report.get("strengths")) if str(x or "").strip()
    ]
    blockers = [
        str(x or "").strip() for x in _as_list(report.get("blockers")) if str(x or "").strip()
    ]
    warnings = [
        str(x or "").strip() for x in _as_list(report.get("warnings")) if str(x or "").strip()
    ]
    recommendations = [
        str(x or "").strip()
        for x in _as_list(report.get("recommendations"))
        if str(x or "").strip()
    ]
    high_conflicts = [
        _as_dict(x)
        for x in _as_list(warning_details.get("high_severity_material_conflicts"))
        if isinstance(x, dict)
    ]
    conflict_recommendations = [
        str(x or "").strip()
        for x in _as_list(warning_details.get("material_conflict_recommendations"))
        if str(x or "").strip()
    ]

    lines = [
        "# 试车前综合体检",
        "",
        f"- 生成时间：`{report.get('generated_at') or '-'}`",
        f"- 项目ID：`{report.get('project_id') or '-'}`",
        f"- 项目名称：`{report.get('project_name') or '-'}`",
        f"- 基础地址：`{report.get('base_url') or '-'}`",
        f"- 试车结论：`{report.get('status_label') or '-'}`",
        f"- 试车是否可执行：`{bool(report.get('trial_run_ready'))}`",
        "",
        "## 签发摘要",
        "",
        f"- 签发决策：`{signoff.get('decision_label') or '-'}`",
        f"- 风险级别：`{signoff.get('risk_label') or '-'}`",
        f"- 签发摘要：`{signoff.get('summary_label') or '-'}`",
        "",
        "## 核验清单",
        "",
    ]
    checklist = _as_list(signoff.get("verification_checklist"))
    if checklist:
        for item in checklist:
            row = _as_dict(item)
            lines.append(
                "- "
                + str(row.get("name") or "-")
                + "：`"
                + ("通过" if bool(row.get("passed")) else "未通过")
                + "`"
                + " / "
                + str(row.get("detail") or "-")
            )
    else:
        lines.append("- 暂无。")

    lines.extend(
        [
            "",
            "## 试车记录草案（待确认）",
            "",
            f"- 记录状态：`{record_draft.get('status_label') or '-'}`",
            f"- 草案摘要：`{record_draft.get('summary_label') or '-'}`",
            f"- 建议结论：`{record_draft.get('recommended_conclusion') or '-'}`",
            f"- 风险级别：`{record_draft.get('recommended_risk_label') or '-'}`",
            f"- 建议试车时间：`{record_draft.get('suggested_executed_at') or '-'}`",
            f"- 执行人：`{record_draft.get('executor_name') or '-'}`",
            f"- 确认提示：`{record_draft.get('confirmation_hint') or '-'}`",
            "",
        ]
    )
    warning_ack_items = [
        str(x or "").strip()
        for x in _as_list(record_draft.get("warning_ack_items"))
        if str(x or "").strip()
    ]
    if warning_ack_items:
        lines.append("### 需确认警告项")
        lines.append("")
        for item in warning_ack_items:
            lines.append(f"- {item}")
        lines.append("")

    next_recommended_action = str(record_draft.get("next_recommended_action") or "").strip()
    if next_recommended_action:
        lines.extend(
            [
                "### 建议优先动作",
                "",
                f"- {next_recommended_action}",
                "",
            ]
        )

    if high_conflicts or conflict_recommendations:
        lines.extend(["## 重点警告明细", ""])
        if high_conflicts:
            lines.extend(["### 高严重度资料冲突清单", ""])
            for item in high_conflicts:
                lines.append(f"- {item.get('summary_label') or '-'}")
                detail_label = str(item.get("detail_label") or "").strip()
                if detail_label:
                    lines.append(f"  - 说明：`{detail_label}`")
                entrypoint_label = str(item.get("entrypoint_label") or "").strip()
                if entrypoint_label:
                    lines.append(f"  - 推荐入口：`{entrypoint_label}`")
                entrypoint_reason_label = str(item.get("entrypoint_reason_label") or "").strip()
                if entrypoint_reason_label:
                    lines.append(f"  - 推荐入口依据：`{entrypoint_reason_label}`")
                review_entrypoint_label = str(
                    item.get("material_review_entrypoint_label") or ""
                ).strip()
                if review_entrypoint_label:
                    lines.append(f"  - 资料核对入口：`{review_entrypoint_label}`")
                material_review_reason_label = str(
                    item.get("material_review_reason_label") or ""
                ).strip()
                if material_review_reason_label:
                    lines.append(f"  - 资料核对依据：`{material_review_reason_label}`")
                action_label = str(item.get("action_label") or "").strip()
                if action_label:
                    lines.append(f"  - 建议动作：`{action_label}`")
                secondary_hint = str(item.get("secondary_hint") or "").strip()
                if secondary_hint:
                    lines.append(f"  - 补充核对：`{secondary_hint}`")
            lines.append("")
        if conflict_recommendations:
            lines.extend(["### 冲突处理建议", ""])
            for item in conflict_recommendations:
                lines.append(f"- {item}")
            lines.append("")

    lines.extend(
        [
            "",
            "## 核心指标",
            "",
            f"- 系统自检：`{bool(metrics.get('self_check_ok'))}`",
            f"- 项目评分前置：`{bool(metrics.get('project_ready_to_score'))}`",
            f"- 项目资料门禁：`{bool(metrics.get('project_gate_passed'))}`",
            f"- 项目 MECE 等级：`{metrics.get('project_mece_level') or '-'}`",
            f"- 真实评标样本：`{metrics.get('ground_truth_count') or 0}`",
            f"- 有效评分记录匹配：`{metrics.get('matched_score_record_count') or 0}`",
            f"- 项目级校准器：`{metrics.get('current_calibrator_version') or '-'}`",
            f"- 演化权重可用：`{bool(metrics.get('evolution_weights_usable'))}`",
            f"- 漂移等级：`{metrics.get('drift_level') or '-'}`",
            f"- 最新评分置信度：`{metrics.get('latest_score_confidence_level') or '-'}`",
            f"- 高严重度资料冲突：`{metrics.get('material_conflict_high_severity_count') or 0}`",
            f"- 系统总封关就绪：`{bool(metrics.get('system_closure_ready'))}`",
            "",
            "## 优势项",
            "",
        ]
    )
    if strengths:
        for item in strengths:
            lines.append(f"- {item}")
    else:
        lines.append("- 暂无。")

    lines.extend(["", "## 阻断项", ""])
    if blockers:
        for item in blockers:
            lines.append(f"- {item}")
    else:
        lines.append("- 无。")

    lines.extend(["", "## 警告项", ""])
    if warnings:
        for item in warnings:
            lines.append(f"- {item}")
    else:
        lines.append("- 无。")

    lines.extend(["", "## 建议动作", ""])
    if recommendations:
        for item in recommendations:
            lines.append(f"- {item}")
    else:
        lines.append("- 当前无需额外动作。")

    return "\n".join(lines).strip() + "\n"


def render_trial_preflight_docx(report: Mapping[str, object]) -> bytes:
    if Document is None:
        raise RuntimeError("DOCX 导出不可用：请安装与当前系统架构兼容的 python-docx/lxml。")

    metrics = _as_dict(report.get("metrics"))
    signoff = _as_dict(report.get("signoff"))
    warning_details = _as_dict(report.get("warning_details"))
    record_draft = _as_dict(report.get("record_draft"))
    strengths = [
        str(x or "").strip() for x in _as_list(report.get("strengths")) if str(x or "").strip()
    ]
    blockers = [
        str(x or "").strip() for x in _as_list(report.get("blockers")) if str(x or "").strip()
    ]
    warnings = [
        str(x or "").strip() for x in _as_list(report.get("warnings")) if str(x or "").strip()
    ]
    recommendations = [
        str(x or "").strip()
        for x in _as_list(report.get("recommendations"))
        if str(x or "").strip()
    ]
    high_conflicts = [
        _as_dict(x)
        for x in _as_list(warning_details.get("high_severity_material_conflicts"))
        if isinstance(x, dict)
    ]
    conflict_recommendations = [
        str(x or "").strip()
        for x in _as_list(warning_details.get("material_conflict_recommendations"))
        if str(x or "").strip()
    ]

    doc = Document()
    doc.add_heading("试车前综合体检", level=0)

    for label, value in (
        ("生成时间", report.get("generated_at") or "-"),
        ("项目ID", report.get("project_id") or "-"),
        ("项目名称", report.get("project_name") or "-"),
        ("基础地址", report.get("base_url") or "-"),
        ("试车结论", report.get("status_label") or "-"),
        ("试车是否可执行", bool(report.get("trial_run_ready"))),
    ):
        doc.add_paragraph(f"{label}：{value}")

    doc.add_heading("签发摘要", level=1)
    for label, value in (
        ("签发决策", signoff.get("decision_label") or "-"),
        ("风险级别", signoff.get("risk_label") or "-"),
        ("签发摘要", signoff.get("summary_label") or "-"),
    ):
        doc.add_paragraph(f"{label}：{value}")

    doc.add_heading("核验清单", level=1)
    checklist = _as_list(signoff.get("verification_checklist"))
    if checklist:
        for item in checklist:
            row = _as_dict(item)
            doc.add_paragraph(
                f"{row.get('name') or '-'}："
                f"{'通过' if bool(row.get('passed')) else '未通过'} / {row.get('detail') or '-'}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("暂无。")

    doc.add_heading("试车记录草案（待确认）", level=1)
    for label, value in (
        ("记录状态", record_draft.get("status_label") or "-"),
        ("草案摘要", record_draft.get("summary_label") or "-"),
        ("建议结论", record_draft.get("recommended_conclusion") or "-"),
        ("风险级别", record_draft.get("recommended_risk_label") or "-"),
        ("建议试车时间", record_draft.get("suggested_executed_at") or "-"),
        ("执行人", record_draft.get("executor_name") or "-"),
        ("确认提示", record_draft.get("confirmation_hint") or "-"),
    ):
        doc.add_paragraph(f"{label}：{value}")
    warning_ack_items = [
        str(x or "").strip()
        for x in _as_list(record_draft.get("warning_ack_items"))
        if str(x or "").strip()
    ]
    if warning_ack_items:
        doc.add_paragraph("需确认警告项：")
        for item in warning_ack_items:
            doc.add_paragraph(item, style="List Bullet")
    next_recommended_action = str(record_draft.get("next_recommended_action") or "").strip()
    if next_recommended_action:
        doc.add_paragraph(f"建议优先动作：{next_recommended_action}")

    if high_conflicts or conflict_recommendations:
        doc.add_heading("重点警告明细", level=1)
        if high_conflicts:
            doc.add_heading("高严重度资料冲突清单", level=2)
            for item in high_conflicts:
                doc.add_paragraph(str(item.get("summary_label") or "-"), style="List Bullet")
                detail_label = str(item.get("detail_label") or "").strip()
                if detail_label:
                    doc.add_paragraph(f"说明：{detail_label}", style="List Bullet 2")
                entrypoint_label = str(item.get("entrypoint_label") or "").strip()
                if entrypoint_label:
                    doc.add_paragraph(f"推荐入口：{entrypoint_label}", style="List Bullet 2")
                entrypoint_reason_label = str(item.get("entrypoint_reason_label") or "").strip()
                if entrypoint_reason_label:
                    doc.add_paragraph(
                        f"推荐入口依据：{entrypoint_reason_label}",
                        style="List Bullet 2",
                    )
                review_entrypoint_label = str(
                    item.get("material_review_entrypoint_label") or ""
                ).strip()
                if review_entrypoint_label:
                    doc.add_paragraph(
                        f"资料核对入口：{review_entrypoint_label}", style="List Bullet 2"
                    )
                material_review_reason_label = str(
                    item.get("material_review_reason_label") or ""
                ).strip()
                if material_review_reason_label:
                    doc.add_paragraph(
                        f"资料核对依据：{material_review_reason_label}",
                        style="List Bullet 2",
                    )
                action_label = str(item.get("action_label") or "").strip()
                if action_label:
                    doc.add_paragraph(f"建议动作：{action_label}", style="List Bullet 2")
                secondary_hint = str(item.get("secondary_hint") or "").strip()
                if secondary_hint:
                    doc.add_paragraph(f"补充核对：{secondary_hint}", style="List Bullet 2")
        if conflict_recommendations:
            doc.add_heading("冲突处理建议", level=2)
            for item in conflict_recommendations:
                doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("核心指标", level=1)
    metric_rows = [
        ("系统自检", bool(metrics.get("self_check_ok"))),
        ("项目评分前置", bool(metrics.get("project_ready_to_score"))),
        ("项目资料门禁", bool(metrics.get("project_gate_passed"))),
        ("项目 MECE 等级", metrics.get("project_mece_level") or "-"),
        ("真实评标样本", metrics.get("ground_truth_count") or 0),
        ("有效评分记录匹配", metrics.get("matched_score_record_count") or 0),
        ("项目级校准器", metrics.get("current_calibrator_version") or "-"),
        ("演化权重可用", bool(metrics.get("evolution_weights_usable"))),
        ("漂移等级", metrics.get("drift_level") or "-"),
        ("最新评分置信度", metrics.get("latest_score_confidence_level") or "-"),
        ("高严重度资料冲突", metrics.get("material_conflict_high_severity_count") or 0),
        ("系统总封关就绪", bool(metrics.get("system_closure_ready"))),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "值"
    for label, value in metric_rows:
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)

    def _append_section(title: str, items: List[str], empty_text: str) -> None:
        doc.add_heading(title, level=1)
        if items:
            for item in items:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph(empty_text)

    _append_section("优势项", strengths, "暂无。")
    _append_section("阻断项", blockers, "无。")
    _append_section("警告项", warnings, "无。")
    _append_section("建议动作", recommendations, "当前无需额外动作。")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
